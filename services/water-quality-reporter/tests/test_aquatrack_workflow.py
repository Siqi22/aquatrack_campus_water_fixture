from io import BytesIO
from pathlib import Path
import tempfile
import unittest
from unittest.mock import patch

from docx import Document
from PIL import Image, ImageDraw

from wqr import (
    Fixture,
    Measurement,
    ReportContext,
    Sample,
    load_profile,
    render_docx,
)


class FakeSupabase:
    configured = True

    def __init__(self):
        self.files = {}
        self.generated_reports = []

    def token(self):
        return "test-token"

    def verify_user(self, _token):
        return {"id": "test-user"}

    def schools(self):
        return [
            {
                "id": "campus-1",
                "name": "Example Elementary",
                "school": "Example Elementary",
                "school_district": "Example District",
                "address": "1 Example Way",
            },
            {
                "id": "campus-2",
                "name": "Example Middle",
                "school": "Example Middle",
                "school_district": "Example District",
                "address": "2 Example Way",
            },
        ]

    def school(self, campus_id):
        return next((school for school in self.schools() if school["id"] == campus_id), None)

    def fixtures(self, campus_id=None):
        if campus_id not in (None, "campus-1", "campus-2"):
            return []
        fixtures = [{
            "id": "fixture-1",
            "campus_id": "campus-1",
            "building_id": "building-1",
            "building_name": "Learning Center",
            "floor": "1",
            "nearest_room": "Hallway A",
            "category": "drinking_fountain",
            "brand": "Example",
            "model": "F100",
            "serial_number": "EX-001",
            "current_result_ppb": 6,
            "current_lead_testing_status": "action_required",
            "current_required_action": "Remediation required",
        }, {
            "id": "fixture-2",
            "campus_id": "campus-2",
            "building_id": "building-2",
            "building_name": "Commons Building",
            "floor": "1",
            "nearest_room": "Main Hall",
            "category": "bottle_filler",
            "brand": "Example",
            "model": "B200",
            "serial_number": "EX-002",
            "current_result_ppb": 3,
            "current_lead_testing_status": "complete",
            "current_required_action": "Complete",
        }]
        return [fixture for fixture in fixtures if campus_id is None or fixture["campus_id"] == campus_id]

    def testing_rounds(self, fixture_ids):
        if "fixture-1" not in fixture_ids:
            return []
        return [{
            "id": "round-1",
            "fixture_id": "fixture-1",
            "round_type": "initial_test",
            "round_number": 1,
            "status": "action_required",
            "sample_id": "SAMPLE-001",
            "sample_draw_date": "2026-03-12",
            "result_value": "6",
            "result_original_unit": "ppb",
            "result_ppb": 6,
            "result_category": "Greater than 5 through 15 ppb",
            "required_action": "Remediation required",
        }]

    def upload_bytes(self, path, data, _content_type="application/octet-stream"):
        self.files[path] = data

    def download_bytes(self, path):
        return self.files.get(path)

    def materialize(self, path, destination):
        data = self.download_bytes(path)
        if data is None:
            return False
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(data)
        return True

    def insert(self, table, row):
        if table == "communication_generated_reports":
            self.generated_reports.append(row)
        return row


def docx_bytes(header_text=None):
    output = BytesIO()
    document = Document()
    if header_text:
        document.sections[0].header.paragraphs[0].text = header_text
    document.add_paragraph("Example communication style")
    document.save(output)
    return output.getvalue()


def pdf_header_bytes():
    output = BytesIO()
    image = Image.new("RGB", (850, 1100), "white")
    drawing = ImageDraw.Draw(image)
    drawing.rectangle((35, 30, 815, 170), fill="#087c92")
    drawing.text((65, 75), "EXAMPLE DISTRICT", fill="white")
    image.save(output, format="PDF", resolution=100)
    return output.getvalue()


class StaticRegistry:
    fixture = Fixture(
        fixture_id="EX-001",
        building="Learning Center",
        floor="1",
        room="Hallway A",
        fixture_type="Drinking fountain",
    )

    def get(self, fixture_id):
        return self.fixture if fixture_id == self.fixture.fixture_id else None


class AquaTrackWorkflowTests(unittest.TestCase):
    def test_database_selection_and_original_report_workflow(self):
        import flask_app
        from supabase_adapter import SupabaseFixtureRegistry

        fake = FakeSupabase()
        with tempfile.TemporaryDirectory() as temp_dir, patch.object(
            flask_app, "supabase", fake
        ), patch.object(
            flask_app, "registry", SupabaseFixtureRegistry(fake)
        ), patch.object(
            flask_app, "WORK", Path(temp_dir)
        ), patch.object(
            flask_app,
            "draft_school_communication_from_reference",
            return_value={
                "layout": "report",
                "intro": "Example introduction.",
                "actions": "Example actions.",
                "notes": "",
                "style_summary": "Example district communication.",
            },
        ):
            client = flask_app.app.test_client()

            start = client.get("/")
            self.assertEqual(start.status_code, 200)
            self.assertIn(b"Example Elementary", start.data)
            self.assertIn(b"<strong>AquaTrack</strong>", start.data)
            self.assertIn(b"Fixture Inventory", start.data)
            self.assertIn(b"Lead Testing", start.data)
            self.assertIn(b"Communication", start.data)
            self.assertIn(b'data-school-search', start.data)
            self.assertIn(b'name="campus_ids"', start.data)
            self.assertIn(b"Example Middle", start.data)
            self.assertIn(b'placeholder="Search by school name"', start.data)
            self.assertNotIn(b"Water Quality Reporter", start.data)
            self.assertNotIn(b">Back<", start.data)
            self.assertNotIn(b"University of Washington", start.data)
            self.assertNotIn(b'value="uw"', start.data)

            setup = client.post(
                "/upload-options",
                data={"campus_ids": ["campus-1", "campus-2"]},
                follow_redirects=False,
            )
            self.assertEqual(setup.status_code, 302)
            self.assertIn("/reference-upload/", setup.headers["Location"])

            fixture_page = client.get(setup.headers["Location"])
            self.assertEqual(fixture_page.status_code, 200)
            self.assertIn(b"Learning Center", fixture_page.data)
            self.assertIn(b"Commons Building", fixture_page.data)
            self.assertIn(b"Example Elementary, Example Middle", fixture_page.data)
            self.assertIn(b"6 ppb", fixture_page.data)
            self.assertIn(b"checked", fixture_page.data)
            self.assertNotIn(b">Back<", fixture_page.data)

            fixture_selection = client.post(
                setup.headers["Location"],
                data={"fixture_ids": ["fixture-1"]},
                follow_redirects=False,
            )
            self.assertIn("/report-style/", fixture_selection.headers["Location"])
            upload_id = fixture_selection.headers["Location"].rstrip("/").split("/")[-1]

            style_page = client.get(fixture_selection.headers["Location"])
            self.assertEqual(style_page.status_code, 200)
            self.assertNotIn(b">Back<", style_page.data)
            self.assertIn(
                b"Upload an example of communication document",
                style_page.data,
            )
            self.assertIn(b"Upload organization header", style_page.data)
            self.assertIn(b"Word document (.docx)", style_page.data)
            self.assertNotIn(b"PDF letterhead", style_page.data)
            self.assertNotIn(b'accept=".docx,.pdf"', style_page.data)
            self.assertNotIn(b"Sample of your report style", style_page.data)

            style = client.post(
                fixture_selection.headers["Location"],
                data={
                    "style_report_file": (
                        BytesIO(pdf_header_bytes()),
                        "style.pdf",
                    ),
                    "header_template_file": (
                        BytesIO(docx_bytes("Example District Header")),
                        "header.docx",
                    ),
                },
                content_type="multipart/form-data",
                follow_redirects=False,
            )
            self.assertIn("/compose/", style.headers["Location"])

            saved_style_page = client.get(fixture_selection.headers["Location"])
            self.assertEqual(saved_style_page.status_code, 200)
            self.assertIn(b"Uploaded file", saved_style_page.data)
            self.assertIn(b"style.pdf", saved_style_page.data)
            self.assertIn(b"header.docx", saved_style_page.data)

            review = client.get(style.headers["Location"])
            self.assertEqual(review.status_code, 200)
            self.assertIn(b"Review and Edit", review.data)
            self.assertIn(b"Sample report style", review.data)
            self.assertIn(b"School District Header", review.data)
            self.assertIn(b"header.docx", review.data)
            self.assertIn(b"source-page-image", review.data)
            self.assertNotIn(b"source-table-wrap", review.data)
            self.assertNotIn(b"page/3.png", review.data)
            self.assertNotIn(b"AquaTrack lead data", review.data)
            self.assertNotIn(b"University of Washington", review.data)
            self.assertNotIn(b">Back<", review.data)

            style_preview = client.get(
                f"/original/{upload_id}/0/page/1.png"
            )
            self.assertEqual(style_preview.status_code, 200)
            self.assertEqual(style_preview.mimetype, "image/png")
            self.assertNotIn(
                "attachment",
                style_preview.headers.get("Content-Disposition", "").lower(),
            )

            generated = client.post(
                style.headers["Location"],
                data={
                    "building": ["Learning Center"],
                    "school_name": "Example Elementary",
                    "organization": "Example District",
                    "sampling_dates": "March 12, 2026",
                    "introduction": "Example introduction.",
                    "actions_taken": "Example actions.",
                    "notes": "",
                    "contact1_name": "Jamie Rivera",
                    "contact1_title": "District Communications",
                    "contact1_phone": "555-0100",
                    "contact1_email": "jamie@example.org",
                },
            )
            self.assertEqual(generated.status_code, 200)
            self.assertEqual(
                generated.mimetype,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            report = Document(BytesIO(generated.data))
            self.assertIn(
                "Example District Header",
                [p.text for p in report.sections[0].header.paragraphs],
            )
            self.assertIn(
                "Table 1",
                "\n".join(paragraph.text for paragraph in report.paragraphs),
            )
            self.assertEqual(len(report.tables), 1)
            results_table = report.tables[0]
            self.assertEqual(
                [cell.text for cell in results_table.rows[0].cells],
                ["#", "Building", "Fixture / Location", "Lead\n(ppb)"],
            )
            self.assertEqual(results_table.rows[1].cells[1].text, "Learning Center")
            self.assertIn("Hallway A", results_table.rows[1].cells[2].text)
            self.assertEqual(results_table.rows[1].cells[3].text, "6")
            report_paragraphs = [paragraph.text for paragraph in report.paragraphs]
            self.assertIn("Questions", report_paragraphs)
            self.assertIn(
                "Jamie Rivera — District Communications — 555-0100, "
                "jamie@example.org",
                report_paragraphs,
            )
            self.assertEqual(len(fake.generated_reports), 1)

    def test_pdf_letterhead_is_accepted_and_embedded_in_word_header(self):
        import flask_app

        header_bytes = pdf_header_bytes()
        flask_app._validate_header_template(header_bytes, "letterhead.pdf")

        with tempfile.TemporaryDirectory() as temp_dir:
            header_path = Path(temp_dir) / "letterhead.pdf"
            header_path.write_bytes(header_bytes)
            sample = Sample(
                sample_id="SAMPLE-001",
                client_sample_id="SAMPLE-001",
                fixture_id="EX-001",
                volume_ml=250,
                collection_date=None,
                analysis_date=None,
                measurements=[Measurement(
                    analyte="Lead",
                    value=6.125,
                    unit="ppb",
                    below_dl=False,
                    detection_limit=None,
                    method="AquaTrack lead testing record",
                )],
            )
            context = ReportContext(
                building="Example Elementary",
                report_date=None,
                sampling_date_range="March 12, 2026",
                introduction_md="Example introduction.",
                actions_taken_md="Example actions.",
                contacts=[],
                samples=[sample],
                action_levels=load_profile("wa_k12_default"),
                analytes_shown=["Lead"],
                report_style="wa_school",
                organization="Example District",
                header_template_path=str(header_path),
            )

            output = BytesIO()
            render_docx(context, StaticRegistry(), output)
            report = Document(BytesIO(output.getvalue()))
            self.assertTrue(
                report.sections[0].header._element.xpath(".//w:drawing")
            )
            self.assertGreater(report.sections[0].top_margin.inches, 1.5)
            self.assertIn(
                "Table 1",
                "\n".join(paragraph.text for paragraph in report.paragraphs),
            )
            self.assertEqual(len(report.tables), 1)


if __name__ == "__main__":
    unittest.main()
