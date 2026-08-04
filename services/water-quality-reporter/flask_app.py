"""Flask app for the Water Quality Reporter.

Multi-user safe: each upload gets a UUID, parsed samples are stored on disk
keyed by that UUID. No global state shared across requests.

Run:  python app.py
"""
from __future__ import annotations
import json
import os
import pickle
import re
import secrets
import uuid
from collections import Counter
from datetime import date, datetime
from io import BytesIO
from pathlib import Path

from flask import (
    Flask, request, render_template, redirect, url_for, send_file,
    flash, abort, jsonify, session, make_response, g,
)
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.utils import secure_filename

from wqr import (
    FixtureRegistry, ReportContext, Measurement, Sample,
    parse_ieh_xlsx, parse_ieh_wide_csv, parse_ieh_pdf,
    parse_generic_lab_file, parse_generic_pdf, parse_doh_school_pdf,
    ClaudePDFConfigurationError, ClaudePDFError,
    parse_school_water_pdf_with_claude, result_pages_from_samples,
    ClaudeStyleError, draft_school_communication_from_reference,
    load_profile, render_docx, evaluate_sample,
)
from wqr.report import _format_sample_volume, _placeholder_fixture
from wqr.parsers import _parse_client_id, _parse_value
from supabase_adapter import SupabaseAdapter, SupabaseFixtureRegistry


HERE = Path(__file__).parent
DATA = HERE / "data"
WORK = Path(os.environ.get("WQR_WORK_DIR", "/tmp/water-quality-reporter")).expanduser()

# Keep the API key server-side. Source runs may use project/.env; the packaged
# preview may use Application Support/Water Quality Reporter Preview/.env.
try:
    from dotenv import load_dotenv

    load_dotenv(HERE / ".env", override=False)
    load_dotenv(WORK.parent / ".env", override=False)
except ImportError:
    pass

WORK.mkdir(parents=True, exist_ok=True)

app = Flask(__name__, template_folder=str(HERE / "flask_templates"))
app.secret_key = (
    os.environ.get("FLASK_SECRET_KEY")
    or os.environ.get("WQR_SECRET_KEY")
    or secrets.token_hex(32)
)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=bool(
        os.environ.get("VERCEL") or os.environ.get("SESSION_COOKIE_SECURE")
    ),
)

supabase = SupabaseAdapter()
registry = (
    SupabaseFixtureRegistry(supabase)
    if supabase.configured
    else FixtureRegistry(DATA / "fixtures.json")
)

DEFAULT_PROFILE_NAME = "wa_k12_default"
REPORT_ANALYTES = ["Lead", "Iron", "Copper", "Manganese", "Zinc"]
NOMENCLATURE_HELP = "Building_Fixture Type_Floor_Room_Volume"
SOURCE_PREVIEW_MAX_ROWS = 50
SOURCE_PREVIEW_MAX_COLS = 14
SOURCE_PREVIEW_MAX_TABLES = 3
SOURCE_PREVIEW_TEXT_CHARS = 1800
AQUATRACK_URL = os.environ.get(
    "AQUATRACK_URL",
    "https://aquatrack-campus-water-fixture.vercel.app",
).rstrip("/")


@app.context_processor
def inject_aquatrack_navigation():
    return {"aquatrack_url": AQUATRACK_URL}


# ---- AquaTrack authentication ----------------------------------------------

@app.before_request
def require_aquatrack_user():
    public_endpoints = {"auth_launch", "auth_session", "health", "preview_health"}
    public_paths = {
        "/health",
        "/auth/launch",
        "/auth/session",
        "/__wqr_preview_health__",
    }
    if request.endpoint in public_endpoints or request.path.rstrip("/") in public_paths:
        return None

    token = supabase.token()
    user = (
        supabase.verify_user(token)
        if supabase.configured
        else {"id": "local-development"}
    )
    if not user:
        return make_response(
            "<!doctype html><title>Sign in required</title>"
            "<body style='font-family:system-ui;max-width:560px;margin:60px auto;padding:20px'>"
            "<h1>Return to AquaTrack</h1>"
            "<p>Your secure reporting session has expired.</p>"
            f"<a href='{os.environ.get('AQUATRACK_URL', '/')}' "
            "style='color:#087c92'>Open AquaTrack</a></body>",
            401,
        )
    g.current_user = user
    session["user_id"] = user.get("id")


@app.get("/health")
def health():
    return {"status": "ok", "supabase_configured": supabase.configured}


@app.get("/auth/launch")
def auth_launch():
    return """<!doctype html><html><head><meta charset="utf-8">
    <meta name="viewport" content="width=device-width,initial-scale=1">
    <title>Opening AquaTrack</title></head>
    <body style="font-family:system-ui;display:grid;place-items:center;min-height:80vh;color:#172333">
    <p id="status">Opening AquaTrack Communication…</p>
    <script>
      const token = new URLSearchParams(location.hash.slice(1)).get('access_token');
      history.replaceState(null, '', location.pathname);
      if (!token) {
        document.getElementById('status').textContent = 'Return to AquaTrack and try again.';
      } else {
        fetch('./session', {
          method: 'POST',
          headers: {'Content-Type': 'application/json'},
          body: JSON.stringify({access_token: token})
        }).then(async response => {
          if (!response.ok) throw new Error(await response.text());
          location.replace('../');
        }).catch(() => {
          document.getElementById('status').textContent =
            'Your AquaTrack session could not be verified.';
        });
      }
    </script></body></html>"""


@app.post("/auth/session")
def auth_session():
    token = (request.get_json(silent=True) or {}).get("access_token", "")
    user = supabase.verify_user(token)
    if not user:
        return {"error": "Invalid or expired AquaTrack session"}, 401
    response = make_response({"ok": True})
    response.set_cookie(
        "wqr_access_token",
        token,
        max_age=3600,
        httponly=True,
        secure=bool(os.environ.get("VERCEL") or request.is_secure),
        samesite="Lax",
    )
    return response


# ---- Building profiles ------------------------------------------------------

def _load_building_profiles() -> dict[str, dict]:
    """Load building_profiles.json. Strip _comment / _README keys."""
    path = DATA / "building_profiles.json"
    if not path.exists():
        return {}
    with open(path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    return {k: v for k, v in raw.items() if not k.startswith("_")}


# Cached at process start. If profiles change, restart the app.
BUILDING_PROFILES = _load_building_profiles()


def _load_building_aliases() -> dict[str, str]:
    """Load uw_building_aliases.json (code -> display name).

    Aliases are pure name mappings. They do NOT carry contact info or any
    other profile data. If a code is in profiles AND aliases, profiles win
    for resolution but aliases stay valid as a name fallback.
    """
    path = DATA / "uw_building_aliases.json"
    if not path.exists():
        return {}
    with open(path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    return {k: v for k, v in raw.items() if not k.startswith("_")}


BUILDING_ALIASES = _load_building_aliases()


def resolve_building(code: str) -> dict:
    """Resolve a building code to a display record. Never raises.

    Lookup order:
      1. Exact match in building_profiles.json (full profile)
      2. Match by code field inside any profile
      3. Match by building name in fixture registry
      4. Match in uw_building_aliases.json (name only)
      5. Fallback: 'Unknown UW Building (CODE)'

    Returns: {
        "code": str,
        "name": str,
        "has_profile": bool,    # full profile entry exists
        "has_alias":   bool,    # name lookup hit aliases (no contacts)
        "has_fixtures": bool,   # any registered fixtures for this building
        "source": "profile" | "alias" | "fixtures" | "unknown",
    }
    """
    code = (code or "").strip().upper()
    if not code:
        return {"code": "", "name": "Unknown Building", "has_profile": False,
                "has_alias": False, "has_fixtures": False, "source": "unknown"}

    profile = BUILDING_PROFILES.get(code)
    if not profile:
        # Profiles can also store the code in a 'code' field
        for p in BUILDING_PROFILES.values():
            if p.get("code", "").upper() == code:
                profile = p
                break

    name_from_alias = BUILDING_ALIASES.get(code)

    # Are there any registered fixtures using this code as a building?
    # Match by either the resolved name or by parsing fixture_id prefix.
    fixture_buildings = {f.building for f in registry.all()}
    fixture_codes = set()
    fixture_name_by_code = {}
    for f in registry.all():
        m = re.match(r"^([A-Z]+\d*)", (f.fixture_id or "").upper())
        if m:
            fixture_code = m.group(1)
            fixture_codes.add(fixture_code)
            fixture_name_by_code.setdefault(fixture_code, f.building)

    if profile:
        return {
            "code": code,
            "name": profile.get("building_name", code),
            "has_profile": True,
            "has_alias": bool(name_from_alias),
            "has_fixtures": code in fixture_codes
                            or profile.get("building_name") in fixture_buildings,
            "source": "profile",
        }

    if name_from_alias:
        return {
            "code": code,
            "name": name_from_alias,
            "has_profile": False,
            "has_alias": True,
            "has_fixtures": code in fixture_codes,
            "source": "alias",
        }

    if code in fixture_name_by_code:
        return {
            "code": code,
            "name": fixture_name_by_code[code],
            "has_profile": False,
            "has_alias": False,
            "has_fixtures": True,
            "source": "fixtures",
        }

    # Try fixture registry by name match (rare — only when fixture's `building`
    # is set to a free-text label that matches the code somehow)
    for fb in fixture_buildings:
        if code.lower() in fb.lower():
            return {
                "code": code,
                "name": fb,
                "has_profile": False,
                "has_alias": False,
                "has_fixtures": True,
                "source": "fixtures",
            }

    return {
        "code": code,
        "name": f"Unknown UW Building ({code})",
        "has_profile": False,
        "has_alias": False,
        "has_fixtures": code in fixture_codes,
        "source": "unknown",
    }


def find_profile(building_name_or_code: str) -> dict:
    """Find a full profile entry. Used for compatibility with code that wants
    the raw contact fields (which only exist in full profiles)."""
    if not building_name_or_code:
        return {}
    needle = building_name_or_code.strip()
    if needle in BUILDING_PROFILES:
        return BUILDING_PROFILES[needle]
    for profile in BUILDING_PROFILES.values():
        if profile.get("building_name", "").strip() == needle:
            return profile
    return {}


def _reported_building_name(sample) -> str:
    """Prefer the report row's building field over inferred school/code names."""
    source_fields = getattr(sample, "source_fields", {}) or {}
    for key in (
        "building_name", "Building Name", "building",
        "facility_name",
    ):
        value = source_fields.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    sample_building = (getattr(sample, "building_name", "") or "").strip()
    if sample_building:
        return sample_building
    for key in ("school_name", "site_name"):
        value = source_fields.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def detected_buildings_from_samples(samples) -> list[dict]:
    """Return each report-provided building, falling back to inferred codes.

        {code, name, has_profile, has_alias, has_fixtures, sample_count}

    Report columns such as Building Name take priority. Results are sorted by
    sample_count descending.
    """
    name_counts = Counter(
        name for s in samples if (name := _reported_building_name(s))
    )
    out = []
    for name, count in name_counts.most_common():
        out.append({
            "code": "",
            "name": name,
            "has_profile": False,
            "has_alias": False,
            "has_fixtures": False,
            "source": "sample",
            "sample_count": count,
        })
    if out:
        return out

    counts = _building_codes_with_counts(samples)
    for code, count in counts:
        rec = resolve_building(code)
        rec["sample_count"] = count
        out.append(rec)
    return out


# ---- Defaults for the compose form ------------------------------------------

def _summarize_samples_for_building(samples: list, building_code: str,
                                     building_name: str) -> dict:
    """Pull the bits the intro paragraph needs from this building's samples.

    Returns:
      fixture_count: int                 # number of unique fixtures
      fixture_phrase: str                # 'metal fountains', 'porcelain fountains',
                                         #   '4 fixtures' etc — pluralized & descriptive
      analytes: list[str]                # in canonical UW order, only those measured
    """
    # Filter samples to this building (by code prefix) so multi-building
    # uploads produce per-building summaries.
    bldg_samples = []
    for s in samples:
        code = _building_code_for(s.fixture_id)
        if (building_code and code == building_code) or (
            not building_code and _reported_building_name(s) == building_name
        ):
            bldg_samples.append(s)
    if not bldg_samples:
        bldg_samples = samples  # fall back to everything

    # Unique fixtures (each fixture often has 250mL + 1L paired samples)
    fixture_ids = {s.fixture_id for s in bldg_samples if s.fixture_id}
    fixture_count = len(fixture_ids)

    # Most common fixture type segment (PF=porcelain, MF=metal, BRS=bottle refill)
    type_codes: list[str] = []
    for fid in fixture_ids:
        parts = fid.split("_")
        for p in parts[1:]:
            if p.upper() in {"PF", "MF", "BRS", "WBF"}:
                type_codes.append(p.upper())
                break
    type_counter = Counter(type_codes)
    type_phrase_singular = {
        "PF":  "porcelain fountain",
        "MF":  "metal fountain",
        "BRS": "bottle refill station",
        "WBF": "water bottle filler",
    }
    if type_counter:
        dominant = type_counter.most_common(1)[0][0]
        singular = type_phrase_singular.get(dominant, "drinking water fixture")
        plural = singular + "s"
    else:
        singular, plural = "drinking water fixture", "drinking water fixtures"

    # Phrase: '4 metal fountains' or 'a metal fountain'
    if fixture_count == 0:
        fixture_phrase = "drinking water fixtures"
    elif fixture_count == 1:
        fixture_phrase = f"one {singular}"
    else:
        # Spell out small numbers, digit for larger
        words = {2: "two", 3: "three", 4: "four", 5: "five",
                 6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten"}
        n = words.get(fixture_count, str(fixture_count))
        fixture_phrase = f"{n} {plural}"

    # Analytes actually measured, in canonical UW report order
    seen_analytes: set = set()
    for s in bldg_samples:
        for m in s.measurements:
            seen_analytes.add(m.analyte)
    canonical_order = REPORT_ANALYTES
    analytes = [a for a in canonical_order if a in seen_analytes]

    return {
        "fixture_count": fixture_count,
        "fixture_phrase": fixture_phrase,
        "analytes": analytes,
    }


def _format_analytes_for_prose(analytes: list[str]) -> str:
    """Format an analyte list as 'iron, lead, copper, manganese, and zinc'."""
    lc = [a.lower() for a in analytes]
    if not lc:
        return "iron, lead, copper, manganese, and zinc"
    if len(lc) == 1:
        return lc[0]
    if len(lc) == 2:
        return f"{lc[0]} and {lc[1]}"
    return ", ".join(lc[:-1]) + ", and " + lc[-1]


# Standard opening paragraph used on every report. Pulled verbatim from
# recurring UW EH&S communications. References the Sieg Building campus-wide
# pilot project — this is intentional context for occupants of any building.
_STANDARD_OPENING = (
    "This message is to inform you that the Environmental Health & Safety "
    "Department (EH&S) and UW Facilities (UWF) have been working to assess "
    "and replace the four porcelain water fountains in the Sieg Building in "
    "coordination with a College of Engineering research group conducting a "
    "pilot research project. You may observe UWF and EH&S flushing these "
    "fixtures, taking additional water samples, placing them out of service, "
    "and replacing plumbing and fixtures."
)

_NEUTRAL_SCHOOL_OPENING = (
    "This message is to share recent drinking water testing results for the "
    "school community. Samples were reviewed using Washington state school "
    "drinking-water guidance, including the 5 ppb lead remediation level in "
    "RCW 28A.210.410."
)


def _default_intro(profile: dict, summary: dict | None = None,
                   building_name: str = "", neutral: bool | None = None) -> str:
    """Auto-generate the introduction paragraph.

    Format matches the Gowen Hall Feb 2026 PDF: standard opening + a
    per-building sentence that names the building, the fixture count, the
    fixture type, and the analytes that were measured.

    `summary` is the dict returned by _summarize_samples_for_building.
    Falls back to a generic sentence when summary is missing (e.g. the
    intro is being computed before samples are parsed).
    """
    name = (profile.get("building_name") or building_name or "the building")
    is_neutral = (
        neutral if neutral is not None
        else not profile and not name.startswith("Unknown UW Building")
    )

    if is_neutral and summary:
        fixture_phrase = summary["fixture_phrase"]
        analytes_prose = _format_analytes_for_prose(summary["analytes"])
        per_building = (
            f"Drinking water outlets at {name} were tested for "
            f"{analytes_prose} using a laboratory report or public health "
            f"results file. This draft is intended to help summarize results, "
            f"identify outlets needing follow-up, and support communication "
            f"with building occupants."
        )
        return _NEUTRAL_SCHOOL_OPENING + "\n\n" + per_building

    if summary:
        fixture_phrase = summary["fixture_phrase"]
        analytes_prose = _format_analytes_for_prose(summary["analytes"])
        per_building = (
            f"The water from {fixture_phrase} in {name} was tested for "
            f"{analytes_prose} using a state-accredited laboratory. The work "
            f"was done in consultation with UWF, EH&S, and the WA Department "
            f"of Health (DOH). The lead testing results are a primary focus "
            f"for public health water quality monitoring."
        )
    else:
        per_building = (
            f"The water from drinking water fixtures in {name} was tested "
            f"for iron, lead, copper, manganese, and zinc using a "
            f"state-accredited laboratory. The work was done in consultation "
            f"with UWF, EH&S, and the WA Department of Health (DOH). The "
            f"lead testing results are a primary focus for public health "
            f"water quality monitoring."
        )

    return _STANDARD_OPENING + "\n\n" + per_building


# ---- Standard UW EH&S contacts ----------------------------------------------
# Per current EH&S workflow these contacts go on every report regardless of
# building. They override any per-building profile contacts. If you need
# different contacts per building, this is the place to make it conditional.

UW_STANDARD_CONTACTS: list[dict] = [
    {
        "name":  "Abebe Aberra",
        "title": "Environmental Public Health Manager",
        "phone": "206-616-1623",
        "email": "aberra@uw.edu",
    },
    {
        "name":  "Dennis Garberg",
        "title": "UW Facilities Associate Director for Maintenance and Construction",
        "phone": "206-221-6501",
        "email": "dgarberg@uw.edu",
    },
]


# ---- Default Actions Taken text --------------------------------------------
# Two templates: the long one when at least one fixture exceeds an action
# level, the short one when everything is clean. Both pulled from the
# recurring language in the UW Art Building communications.

ACTIONS_WITH_EXCEEDANCES = (
    "**1. Flushing Fixtures Thoroughly:** The more time water has been sitting "
    "in pipes, the more lead and other metals it may contain. UW Facilities "
    "may flush drinking water fixtures to reduce stagnation as part of "
    "follow-up work.\n\n"
    "**2. Shutting Off Water Fixtures of Concern:** Fixtures exceeding the "
    "default review level should be posted out of service or made "
    "inaccessible while follow-up testing, replacement, or remediation is "
    "completed.\n\n"
    "**3. Providing Alternative Drinking Water:** Where fixtures are taken "
    "out of service, building occupants should continue to have access to "
    "alternative drinking water until fixtures are returned to service or "
    "replaced.\n\n"
    "**4. Replacement of Water Fountains and Associated Plumbing:** As part "
    "of a campus-wide effort to upgrade water infrastructure, UW Facilities "
    "may replace older water fountains with bottle refill stations. New "
    "fountains should include NSF/ANSI-certified lead-free parts and "
    "filtration designed to capture metals and other contaminants. The "
    "fountains and filters will be serviced and maintained by UW Facilities "
    "on an ongoing basis.\n\n"
    "**5. Conducting Follow-Up Evaluation:** UW EH&S, in coordination with "
    "UW Facilities, will continue to evaluate potential lead sources, "
    "including fixture components, plumbing components, or premise plumbing, "
    "as needed."
)

NEUTRAL_ACTIONS_WITH_EXCEEDANCES = (
    "**1. Results 0-5 ppb:** No immediate action is needed for outlets with "
    "lead results of 5 ppb or less. Schools built, or with all plumbing "
    "replaced, before 2016 should plan for required retesting at least once "
    "every five years beginning July 1, 2026.\n\n"
    "**2. Results greater than 5 ppb:** Outlets above 5 ppb require "
    "remediation. Schools with one or more outlets above 5 ppb should develop "
    "an action plan in consultation with DOH or the local health jurisdiction "
    "and communicate annually with the school community.\n\n"
    "**3. Results greater than 15 ppb:** Schools should immediately shut off "
    "water to the outlet or make the outlet inaccessible for consumption, "
    "such as by bagging or taping off the area and posting signage.\n\n"
    "**4. After remediation:** Finish the fixture replacement process by "
    "conditioning the fixture(s) and conducting follow-up sampling/testing to "
    "confirm results are 5 ppb or lower.\n\n"
    "If you need help communicating this with your building occupants, please "
    "let us know."
)

ACTIONS_NO_EXCEEDANCES = (
    "**1. Continued Monitoring:** EH&S and UW Facilities will continue to "
    "review sampling results and maintain records for the tested fixtures.\n\n"
    "**2. Infrastructure Planning:** UW Facilities may continue planning "
    "fixture upgrades as part of broader campus water infrastructure "
    "improvements."
)


def _default_actions(profile: dict, has_exceedances: bool = True,
                     neutral: bool = False) -> str:
    """Choose actions template based on whether any fixture exceeded a level."""
    if neutral and has_exceedances:
        return NEUTRAL_ACTIONS_WITH_EXCEEDANCES
    return ACTIONS_WITH_EXCEEDANCES if has_exceedances else ACTIONS_NO_EXCEEDANCES


def _default_contacts(profile: dict, neutral: bool = False) -> list[dict]:
    """Standard UW EH&S contacts go on every report.

    `profile` is accepted but ignored — kept in the signature so the rest of
    the codebase doesn't need to change. If per-building contact overrides
    become a real need, this is where to add the lookup.
    """
    if neutral:
        return [
            {"name": "", "title": "", "phone": "", "email": ""},
            {"name": "", "title": "", "phone": "", "email": ""},
        ]
    return [dict(c) for c in UW_STANDARD_CONTACTS]


def defaults_for(building_name: str, has_exceedances: bool = True,
                 samples: list | None = None, report_style: str | None = None,
                 display_name: str = "") -> dict:
    """Return the bundle of default text + contacts for one building.

    has_exceedances drives the Actions Taken template choice (long vs. short).
    samples (the parsed lab samples for this upload) are used to produce a
    data-driven intro paragraph naming fixture count, fixture type, and
    analytes for the chosen building.
    """
    profile = find_profile(building_name)
    summary = None
    if samples is not None:
        # Find this building's code so we can summarize only its samples
        code = profile.get("code", "")
        if not code:
            # Reverse-resolve: look up code from name
            for c, name in BUILDING_ALIASES.items():
                if name == building_name:
                    code = c
                    break
        if not code and "(" in building_name and building_name.endswith(")"):
            # 'Unknown UW Building (XXX)' — pull the code from the parens
            code = building_name.rsplit("(", 1)[1].rstrip(")")
        summary = _summarize_samples_for_building(
            samples, code.upper() if code else "", building_name)
    if report_style == "uw":
        neutral = False
    elif report_style == "wa_school":
        neutral = True
    else:
        neutral = not profile and not building_name.startswith("Unknown UW Building")

    intro_name = display_name.strip() or building_name
    return {
        "intro":    _default_intro(profile, summary, intro_name, neutral=neutral),
        "actions":  _default_actions(
            profile, has_exceedances,
            neutral=neutral,
        ),
        "contacts": _default_contacts(profile, neutral=neutral),
    }


# ---- Sample-to-building inference -------------------------------------------

# Match the leading alphanumeric "code" portion of a fixture_id.
# Letters followed optionally by digits, then an underscore.
_SAMPLE_PREFIX_RE = re.compile(r"^([A-Z]+\d*)_")
_LETTERS_ONLY_RE = re.compile(r"^([A-Z]+)")


def _building_code_for(fixture_id: str) -> str | None:
    """Extract the building code from a fixture_id.

    Tries the literal prefix first (e.g. 'ART_PF_025' -> 'ART').
    If that prefix isn't a known building, strips trailing digits and tries
    again — this handles legacy IDs like 'SIG1_PF_200' where 'SIG1' means
    'Sieg sample #1' not the building code.
    """
    if not fixture_id:
        return None
    m = _SAMPLE_PREFIX_RE.match(fixture_id)
    if not m:
        return None
    raw = m.group(1)
    # Try literal match first
    if raw in BUILDING_PROFILES:
        return raw
    if raw in BUILDING_ALIASES:
        return raw
    # Strip trailing digits (legacy formats: SIG1, SIG2, ...)
    m2 = _LETTERS_ONLY_RE.match(raw)
    if m2 and m2.group(1) in BUILDING_PROFILES:
        return m2.group(1)
    if m2 and m2.group(1) in BUILDING_ALIASES:
        return m2.group(1)
    # Final fallback: return raw prefix even though it's not in profiles,
    # so the UI can flag it as unknown.
    return raw


def _building_codes_with_counts(samples) -> list[tuple[str, int]]:
    """All building codes in sample IDs, with counts. Sorted by count desc."""
    codes = [c for c in (_building_code_for(s.fixture_id) for s in samples) if c]
    return Counter(codes).most_common() if codes else []


def _infer_building_code_from_samples(samples) -> str | None:
    """Most common building code prefix in sample IDs."""
    counts = _building_codes_with_counts(samples)
    return counts[0][0] if counts else None


def _suggested_building_for(samples) -> str:
    """Best guess of which building this upload is for."""
    reported = Counter(
        name for s in samples if (name := _reported_building_name(s))
    )
    if reported:
        return reported.most_common(1)[0][0]
    code = _infer_building_code_from_samples(samples)
    if code:
        return resolve_building(code)["name"]
    return ""


# ---- Upload helpers ---------------------------------------------------------

def _parse_upload(filename: str, raw_bytes: bytes):
    """Dispatch by extension. Returns list[Sample] or raises ValueError."""
    suffix = Path(filename).suffix.lower()
    tmp = WORK / f"upload_{uuid.uuid4().hex}{suffix}"
    tmp.write_bytes(raw_bytes)
    try:
        if suffix == ".pdf":
            errors = []
            for parser in (parse_ieh_pdf, parse_doh_school_pdf, parse_generic_pdf):
                try:
                    samples = parser(tmp)
                    if samples:
                        return samples
                except Exception as e:
                    errors.append(str(e))
            raise ValueError(
                "Could not confidently parse this PDF. IEH PDFs must include "
                "a recognizable Certificate of Analysis table. WA DOH school "
                "reports must include the standard results table. For other "
                "lab outputs, please upload CSV or XLSX."
            )
        elif suffix in {".xlsx", ".xls"}:
            parsers = [parse_ieh_xlsx, parse_generic_lab_file]
        elif suffix == ".csv":
            parsers = [parse_ieh_wide_csv, parse_generic_lab_file]
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        errors = []
        for parser in parsers:
            try:
                samples = parser(tmp)
                if samples:
                    return samples
            except Exception as e:
                errors.append(str(e))
        detail = errors[-1] if errors else "No samples found."
        raise ValueError(detail)
    finally:
        tmp.unlink(missing_ok=True)


def _parse_school_testing_result(filename: str, raw_bytes: bytes):
    """Parse a school result file into the canonical fixture-result schema.

    PDFs use Claude's format-agnostic visual extraction first. The older local
    adapters remain an offline/service-failure fallback for known files only.
    CSV and Excel inputs continue through the generic tabular parser.
    Returns ``(samples, parser_source, fixture_result_pages)``.
    """
    suffix = Path(filename).suffix.lower()
    if suffix != ".pdf":
        return _parse_upload(filename, raw_bytes), "local", []

    claude_error: ClaudePDFError | None = None
    try:
        samples = parse_school_water_pdf_with_claude(
            raw_bytes,
            filename,
        )
        return (
            samples,
            "claude-haiku-4-5-20251001",
            result_pages_from_samples(samples),
        )
    except ClaudePDFError as exc:
        claude_error = exc

    try:
        samples = _parse_upload(filename, raw_bytes)
    except ValueError as local_error:
        if isinstance(claude_error, ClaudePDFConfigurationError):
            raise ValueError(
                "This PDF does not match a known offline format, and "
                "CLAUDE_API_KEY is not configured for format-agnostic parsing."
            ) from local_error
        raise ClaudePDFError(
            "Claude could not extract this PDF, and the known-format fallback "
            "could not parse it either."
        ) from local_error
    return samples, "local-fallback", _school_result_pages_from_bytes(raw_bytes)


def _save_session_data(upload_id: str, samples, meta: dict):
    blob = WORK / f"session_{upload_id}.pkl"
    payload = pickle.dumps({"samples": samples, "meta": meta})
    blob.write_bytes(payload)
    if supabase.configured:
        supabase.upload_bytes(
            f"{g.current_user['id']}/sessions/{upload_id}.pkl",
            payload,
        )


def _load_session_data(upload_id: str):
    blob = WORK / f"session_{upload_id}.pkl"
    if not blob.exists() and supabase.configured:
        supabase.materialize(
            f"{g.current_user['id']}/sessions/{upload_id}.pkl",
            blob,
        )
    if not blob.exists():
        return None
    data = pickle.loads(blob.read_bytes())
    for entry in data.get("meta", {}).get("original_files", []):
        path = Path(entry.get("path", ""))
        storage_path = entry.get("storage_path")
        if storage_path and not path.exists():
            supabase.materialize(storage_path, path)
    return data


def _original_file_for(meta: dict, field: str) -> dict | None:
    return next(
        (
            entry
            for entry in meta.get("original_files", [])
            if entry.get("field") == field
        ),
        None,
    )


def _session_has_style(meta: dict) -> bool:
    return _original_file_for(meta, "style_report_file") is not None


def _session_has_header(meta: dict) -> bool:
    return _original_file_for(meta, "header_template_file") is not None


def _session_can_review(meta: dict) -> bool:
    return (
        _session_has_style(meta)
        and _session_has_header(meta)
        and not meta.get("style_needs_refresh", False)
    )


def _clean_source_cell(value) -> str:
    text = "" if value is None else str(value)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:120] + "..." if len(text) > 120 else text


def _dataframe_source_table(df, title: str) -> dict | None:
    if df is None or df.empty:
        return None
    df = df.fillna("")
    df = df.iloc[:SOURCE_PREVIEW_MAX_ROWS, :SOURCE_PREVIEW_MAX_COLS]
    headers = [_clean_source_cell(c) or f"Column {i + 1}" for i, c in enumerate(df.columns)]
    rows = [
        [_clean_source_cell(value) for value in row]
        for row in df.astype(str).values.tolist()
    ]
    if not any(any(cell for cell in row) for row in rows):
        return None
    return {
        "title": title,
        "headers": headers,
        "rows": rows,
        "truncated": len(df.index) >= SOURCE_PREVIEW_MAX_ROWS,
    }


def _raw_source_table(raw_rows, title: str) -> dict | None:
    cleaned = []
    width = 0
    for row in raw_rows or []:
        cells = [_clean_source_cell(value) for value in (row or [])]
        if not any(cells):
            continue
        cleaned.append(cells)
        width = max(width, len(cells))
    if not cleaned or width < 2:
        return None
    width = min(width, SOURCE_PREVIEW_MAX_COLS)
    rows = [
        (row + [""] * width)[:width]
        for row in cleaned[:SOURCE_PREVIEW_MAX_ROWS]
    ]
    return {
        "title": title,
        "headers": [f"Col {i + 1}" for i in range(width)],
        "rows": rows,
        "truncated": len(cleaned) > SOURCE_PREVIEW_MAX_ROWS,
    }


def _school_result_page_numbers(pdf) -> list[int]:
    """Return 1-based pages that contain fixture-level result rows.

    DOH reports may begin with explanations and definitions. A result page
    must contain sample/fixture/result language plus either a source-style
    numeric sample ID or a non-empty extracted table. This detector is used
    only by the known-format fallback and by display previews; it never limits
    the pages sent to Claude. An empty list means no confident local match.
    """
    detected: list[int] = []
    for page_number, page in enumerate(pdf.pages, start=1):
        text = page.extract_text() or ""
        lowered = text.lower()
        has_result_vocabulary = (
            "sample" in lowered
            and "fixture" in lowered
            and "result" in lowered
        )
        if not has_result_vocabulary:
            continue
        has_sample_ids = bool(re.search(r"(?<!\d)\d{6}(?!\d)", text))
        has_table_rows = any(
            table and len(table) >= 2
            for table in (page.extract_tables() or [])
        )
        if has_sample_ids or has_table_rows:
            detected.append(page_number)
    return detected


def _school_result_pages_from_bytes(raw_bytes: bytes) -> list[int]:
    import pdfplumber

    with pdfplumber.open(BytesIO(raw_bytes)) as pdf:
        return _school_result_page_numbers(pdf)


def _build_source_preview(
    path: Path,
    filename: str,
    *,
    result_pages_only: bool = False,
    relevant_pages_override: list[int] | None = None,
) -> dict:
    """Small source-file preview for compose-page QC.

    This is deliberately display-only. Parsing confidence comes from the
    canonical extraction/validation pipeline; the preview helps the user
    compare the original source to the editable parsed table without opening
    another tab.
    """
    suffix = path.suffix.lower()
    preview = {
        "kind": suffix.lstrip(".") or "file",
        "is_pdf": suffix == ".pdf",
        "tables": [],
        "text": "",
        "error": "",
        "relevant_pages": [],
        "omitted_pages": 0,
    }
    try:
        if suffix == ".csv":
            import pandas as pd

            df = pd.read_csv(path, dtype=str, keep_default_na=False, nrows=SOURCE_PREVIEW_MAX_ROWS)
            table = _dataframe_source_table(df, filename)
            if table:
                preview["tables"].append(table)
            return preview

        if suffix in {".xls", ".xlsx"}:
            import pandas as pd

            df = pd.read_excel(path, dtype=str, keep_default_na=False, nrows=SOURCE_PREVIEW_MAX_ROWS)
            table = _dataframe_source_table(df, filename)
            if table:
                preview["tables"].append(table)
            return preview

        if suffix == ".docx":
            from docx import Document

            doc = Document(str(path))
            header_paragraphs: list[str] = []
            seen_header_parts: set[str] = set()
            for section_index, section in enumerate(doc.sections, start=1):
                header = section.header
                part_name = str(header.part.partname)
                if part_name in seen_header_parts:
                    continue
                seen_header_parts.add(part_name)
                header_paragraphs.extend(
                    p.text.strip() for p in header.paragraphs if p.text.strip()
                )
                for table_index, header_table in enumerate(
                    header.tables[:SOURCE_PREVIEW_MAX_TABLES], start=1
                ):
                    raw_table = [
                        [_clean_source_cell(cell.text) for cell in row.cells]
                        for row in header_table.rows[:SOURCE_PREVIEW_MAX_ROWS]
                    ]
                    table = _raw_source_table(
                        raw_table,
                        f"Header table {table_index}",
                    )
                    if table:
                        preview["tables"].append(table)

            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text_blocks = []
            if header_paragraphs:
                text_blocks.append("HEADER\n" + "\n".join(header_paragraphs))
            if paragraphs:
                text_blocks.append("DOCUMENT BODY\n" + "\n".join(paragraphs))
            preview["text"] = "\n\n".join(text_blocks)[:SOURCE_PREVIEW_TEXT_CHARS]
            for table_index, doc_table in enumerate(doc.tables[:SOURCE_PREVIEW_MAX_TABLES], start=1):
                raw_table = [
                    [_clean_source_cell(cell.text) for cell in row.cells]
                    for row in doc_table.rows[:SOURCE_PREVIEW_MAX_ROWS]
                ]
                table = _raw_source_table(raw_table, f"Word table {table_index}")
                if table:
                    preview["tables"].append(table)
            if not preview["text"] and not preview["tables"]:
                preview["text"] = "No readable text or tables were found in this Word file."
            return preview

        if suffix == ".pdf":
            import pdfplumber

            text_parts: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                all_pages = list(range(1, len(pdf.pages) + 1))
                if result_pages_only:
                    detected_pages = (
                        relevant_pages_override
                        if relevant_pages_override is not None
                        else _school_result_page_numbers(pdf)
                    )
                    relevant_pages = [
                        page for page in detected_pages if page in all_pages
                    ] or all_pages
                else:
                    relevant_pages = all_pages
                preview["relevant_pages"] = relevant_pages
                preview["omitted_pages"] = len(pdf.pages) - len(relevant_pages)
                for page_index, page in enumerate(pdf.pages):
                    page_number = page_index + 1
                    if page_number not in relevant_pages:
                        continue
                    if len(" ".join(text_parts)) < SOURCE_PREVIEW_TEXT_CHARS:
                        page_text = page.extract_text() or ""
                        if page_text:
                            text_parts.append(page_text)

                    if len(preview["tables"]) >= SOURCE_PREVIEW_MAX_TABLES:
                        continue
                    for table_index, raw_table in enumerate(page.extract_tables() or [], start=1):
                        table = _raw_source_table(
                            raw_table,
                            f"Page {page_index + 1} table {table_index}",
                        )
                        if table:
                            preview["tables"].append(table)
                        if len(preview["tables"]) >= SOURCE_PREVIEW_MAX_TABLES:
                            break
            text = re.sub(r"\s+\n", "\n", "\n\n".join(text_parts)).strip()
            preview["text"] = text[:SOURCE_PREVIEW_TEXT_CHARS]
            return preview

        preview["text"] = "Inline preview is not available for this file type. Use the source-file link to review it."
        return preview
    except Exception as e:
        preview["error"] = f"Could not build an inline preview for this file: {e}"
        return preview


def _detect_reference_headings(text: str) -> list[str]:
    heading_words = {
        "introduction", "background", "summary", "results", "findings",
        "actions", "action", "recommendations", "testing", "sampling",
        "health", "contacts", "contact", "next steps", "remediation",
        "communication", "limitations",
    }
    headings: list[str] = []
    for raw in (text or "").splitlines():
        line = re.sub(r"\s+", " ", raw).strip(" :-")
        if not line or len(line) > 70:
            continue
        lowered = line.lower()
        if lowered in heading_words or any(word in lowered for word in heading_words):
            if line not in headings:
                headings.append(line)
        if len(headings) >= 6:
            break
    return headings


def _reference_notes_from_files(files: list[dict]) -> list[str]:
    notes: list[str] = []
    style_file = next((f for f in files if f.get("field") == "style_report_file"), None)
    lab_file = next((
        f for f in files
        if f.get("field") in {"school_testing_result_file", "reference_lab_file"}
    ), None)

    if style_file:
        preview = style_file.get("preview", {})
        text = preview.get("text", "")
        headings = _detect_reference_headings(text)
        if headings:
            notes.append("Reference report sections detected: " + "; ".join(headings) + ".")
        if re.search(r"\b(Date|To|From|Subject)\s*:", text, flags=re.I):
            notes.append("Reference report appears to use a memo-style header.")
        if preview.get("tables"):
            notes.append(f"Reference report includes {len(preview['tables'])} readable table preview(s).")
        if text and not headings:
            notes.append("Reference report text was extracted for review in the QC panel.")

    if lab_file:
        preview = lab_file.get("preview", {})
        tables = preview.get("tables", [])
        if tables:
            headers = [h for h in tables[0].get("headers", []) if h]
            if headers:
                notes.append("School testing-result columns include: " + ", ".join(headers[:8]) + ".")
        elif preview.get("text"):
            notes.append("School testing-result text was extracted for review in the QC panel.")

    if style_file or lab_file:
        notes.append("Use these notes and the QC panel to adjust the editable narrative before generating the DOCX.")
    return notes


def _communication_facts(samples, setup: dict) -> dict:
    """Compact, deterministic facts for reference-style drafting.

    The style model never receives the testing-result PDF. Measurements and
    classifications are calculated locally and supplied as locked facts.
    """
    lead_rows = []
    fixture_ids = set()
    analytes = set()
    building_names = set()
    collection_dates = set()
    for sample in samples:
        fixture_ids.add(sample.fixture_id or sample.sample_id)
        if sample.collection_date:
            collection_dates.add(sample.collection_date)
        analytes.update(m.analyte for m in sample.measurements)
        reported_building = _reported_building_name(sample)
        if reported_building:
            building_names.add(reported_building)
        lead = sample.measurement("Lead")
        if lead is None:
            continue
        lead_rows.append({
            "sample_id": sample.client_sample_id or sample.sample_id,
            "fixture_location": sample.fixture_label or "not provided",
            "building_name": reported_building or "not provided",
            "reported_lead_ppb": lead.display,
            "numeric_lead_ppb": (
                lead.value if not lead.below_dl and lead.value is not None else None
            ),
        })

    elevated = [
        row for row in lead_rows
        if row["numeric_lead_ppb"] is not None and row["numeric_lead_ppb"] > 5
    ]
    immediate = [
        row for row in lead_rows
        if row["numeric_lead_ppb"] is not None and row["numeric_lead_ppb"] > 15
    ]
    ordered_dates = sorted(collection_dates)
    if not ordered_dates:
        collection_date_range = "not provided"
    elif ordered_dates[0] == ordered_dates[-1]:
        collection_date_range = _format_date(ordered_dates[0])
    else:
        collection_date_range = (
            f"{_format_date(ordered_dates[0])} through "
            f"{_format_date(ordered_dates[-1])}"
        )
    return {
        "school_name": setup.get("school_name") or "[confirm school name]",
        "organization": setup.get("organization") or "[confirm district or organization]",
        "collection_date_range": collection_date_range,
        "sample_count": len(samples),
        "tested_fixture_or_outlet_count": len(fixture_ids),
        "analytes_measured": sorted(analytes),
        "building_names": sorted(building_names),
        "lead_remediation_threshold_ppb": 5,
        "lead_immediate_response_threshold_ppb": 15,
        "outlets_above_5_ppb": len(elevated),
        "outlets_above_15_ppb": len(immediate),
        "maximum_detected_lead_ppb": max(
            (row["numeric_lead_ppb"] for row in lead_rows
             if row["numeric_lead_ppb"] is not None),
            default=None,
        ),
        "elevated_outlets": elevated[:30],
        "confirmed_actions": [],
        "missing_operational_details": [
            "whether affected outlets were shut off or made inaccessible",
            "remediation status and timeline",
            "district website and contact information",
        ],
    }


SOURCE_FILE_LABELS = {
    "parsed": "Current lab/results file",
    "coc_file": "Original COC / sampling form",
    "style_report_file": "Sample report style",
    "header_template_file": "School District Header",
    "school_testing_result_file": "School Water Testing Result",
    "reference_lab_file": "School Water Testing Result",
}


def _validate_header_template(raw_bytes: bytes, filename: str) -> None:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        try:
            import pypdfium2 as pdfium

            document = pdfium.PdfDocument(raw_bytes)
            try:
                if len(document) < 1:
                    raise ValueError("The PDF does not contain any pages.")
                page = document[0]
                try:
                    width, height = page.get_size()
                finally:
                    page.close()
            finally:
                document.close()
            if width <= 0 or height <= 0:
                raise ValueError("The first PDF page has an invalid size.")
            return
        except Exception as exc:
            raise ValueError(
                "The Header template could not be read as a PDF file."
            ) from exc

    if suffix != ".docx":
        raise ValueError("The Header template must be a DOCX or PDF file.")
    try:
        from docx import Document

        document = Document(BytesIO(raw_bytes))
    except Exception as exc:
        raise ValueError("The Header template could not be read as a DOCX file.") from exc

    has_header_content = False
    seen_parts: set[str] = set()
    for section in document.sections:
        header = section.header
        part_name = str(header.part.partname)
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        if any(p.text.strip() for p in header.paragraphs):
            has_header_content = True
        if any(
            cell.text.strip()
            for table in header.tables
            for row in table.rows
            for cell in row.cells
        ):
            has_header_content = True
        if header._element.xpath(".//w:drawing | .//w:pict"):
            has_header_content = True
    if not has_header_content:
        raise ValueError(
            "The Header template must contain content in the Word header area."
        )


def _source_file_entry(
    field: str,
    filename: str,
    path: Path,
    *,
    relevant_pages: list[int] | None = None,
) -> dict:
    entry = {
        "field": field,
        "label": SOURCE_FILE_LABELS.get(field, field.replace("_", " ").title()),
        "filename": filename,
        "path": str(path),
        "preview": _build_source_preview(
            path,
            filename,
            result_pages_only=field == "school_testing_result_file",
            relevant_pages_override=relevant_pages,
        ),
    }
    if relevant_pages:
        entry["relevant_pages"] = list(relevant_pages)
    if supabase.configured:
        storage_path = f"{g.current_user['id']}/files/{path.parent.name}/{path.name}"
        supabase.upload_bytes(storage_path, path.read_bytes())
        entry["storage_path"] = storage_path
    return entry


def _save_original_upload(upload_id: str, field: str, f) -> dict | None:
    if not f or not f.filename:
        return None
    safe = secure_filename(f.filename) or f"{field}.dat"
    folder = WORK / f"originals_{upload_id}"
    folder.mkdir(exist_ok=True)
    path = folder / f"{field}_{safe}"
    f.stream.seek(0)
    path.write_bytes(f.read())
    return _source_file_entry(field, f.filename, path)


def _save_reference_upload(setup_id: str, field: str, f) -> dict | None:
    if not f or not f.filename:
        return None
    safe = secure_filename(f.filename) or f"{field}.dat"
    folder = WORK / f"references_{setup_id}"
    folder.mkdir(exist_ok=True)
    path = folder / f"{field}_{safe}"
    f.stream.seek(0)
    path.write_bytes(f.read())
    return _source_file_entry(field, f.filename, path)


def _save_reference_setup(setup_id: str, setup: dict, reference_files: list[dict]) -> None:
    blob = WORK / f"setup_{setup_id}.pkl"
    payload = pickle.dumps({"setup": setup, "reference_files": reference_files})
    blob.write_bytes(payload)
    if supabase.configured:
        supabase.upload_bytes(
            f"{g.current_user['id']}/setups/{setup_id}.pkl",
            payload,
        )


def _load_reference_setup(setup_id: str) -> dict | None:
    blob = WORK / f"setup_{setup_id}.pkl"
    if not blob.exists() and supabase.configured:
        supabase.materialize(
            f"{g.current_user['id']}/setups/{setup_id}.pkl",
            blob,
        )
    if not blob.exists():
        return None
    data = pickle.loads(blob.read_bytes())
    for entry in data.get("reference_files", []):
        path = Path(entry.get("path", ""))
        storage_path = entry.get("storage_path")
        if storage_path and not path.exists():
            supabase.materialize(storage_path, path)
    return data


def _report_setup_from_values(values) -> dict:
    requested_ids = values.getlist("campus_ids") if hasattr(values, "getlist") else []
    legacy_id = (values.get("campus_id") or "").strip()
    if not requested_ids and legacy_id:
        requested_ids = [legacy_id]
    requested_ids = list(dict.fromkeys(item.strip() for item in requested_ids if item.strip()))
    selected_schools = [
        school
        for campus_id in requested_ids
        if (school := supabase.school(campus_id)) is not None
    ] if supabase.configured else []
    campus_ids = [school["id"] for school in selected_schools]
    school_names = [
        school.get("school") or school.get("name") or "School"
        for school in selected_schools
    ]
    organization = selected_schools[0].get("school_district") if selected_schools else ""

    return {
        "report_style": "wa_school",
        "organization": organization,
        "school_name": ", ".join(school_names),
        "school_names": school_names,
        "campus_id": campus_ids[0] if campus_ids else "",
        "campus_ids": campus_ids,
    }


def _school_options() -> list[dict]:
    return supabase.schools() if supabase.configured else []


def _fixtures_for_campuses(campus_ids: list[str]) -> list[dict]:
    school_by_id = {school["id"]: school for school in _school_options()}
    fixtures = []
    seen = set()
    for campus_id in campus_ids:
        school = school_by_id.get(campus_id) or {}
        school_name = school.get("school") or school.get("name") or "School"
        for fixture in supabase.fixtures(campus_id):
            if fixture["id"] in seen:
                continue
            seen.add(fixture["id"])
            fixtures.append({**fixture, "school_name": school_name})
    return fixtures


def _samples_from_inventory(campus_ids: list[str], fixture_ids: list[str]) -> tuple[list, list[dict]]:
    allowed_fixtures = _fixtures_for_campuses(campus_ids)
    selected_ids = set(fixture_ids)
    selected = [row for row in allowed_fixtures if row["id"] in selected_ids]
    if len(selected) != len(selected_ids):
        abort(403)

    rounds = supabase.testing_rounds([row["id"] for row in selected])
    latest_round = {}
    for testing_round in rounds:
        latest_round.setdefault(testing_round["fixture_id"], testing_round)

    samples = []
    for fixture in selected:
        testing_round = latest_round.get(fixture["id"]) or {}
        registry_id = fixture.get("serial_number") or fixture["id"]
        result_ppb = testing_round.get("result_ppb")
        if result_ppb is None:
            result_ppb = fixture.get("current_result_ppb")

        measurements = []
        if result_ppb is not None:
            measurements.append(Measurement(
                analyte="Lead",
                value=float(result_ppb),
                unit="ppb",
                below_dl=False,
                detection_limit=None,
                method="AquaTrack lead testing record",
            ))

        sample_date_text = testing_round.get("sample_draw_date") or ""
        sample_date = date.fromisoformat(sample_date_text) if sample_date_text else None
        sample_id = (
            testing_round.get("sample_id")
            or fixture.get("serial_number")
            or fixture["id"]
        )
        location = (
            f"{fixture.get('building_name') or 'Building'} · "
            f"Floor {fixture.get('floor') or '—'} · "
            f"{fixture.get('nearest_room') or 'Location not recorded'}"
        )
        samples.append(Sample(
            sample_id=sample_id,
            client_sample_id=sample_id,
            fixture_id=registry_id,
            volume_ml=250,
            collection_date=sample_date,
            analysis_date=None,
            measurements=measurements,
            building_name=fixture.get("building_name") or "",
            fixture_label=location,
            source_fields={
                "aquatrack_fixture_id": fixture["id"],
                "lead_testing_round_id": testing_round.get("id"),
                "lead_testing_status": (
                    fixture.get("current_lead_testing_status") or "not_started"
                ),
                "school_name": fixture.get("school_name") or "",
            },
        ))
    return samples, selected


# ---- Routes ----------------------------------------------------------------

@app.route("/")
def index():
    return render_template(
        "start.html",
        upload_id="",
        campus_id="",
        campus_ids=[],
        school_name="",
        schools=_school_options(),
        aquatrack_url=os.environ.get("AQUATRACK_URL", "/"),
    )


@app.get("/__wqr_preview_health__")
def preview_health():
    return "water-quality-reporter-preview:2026-07-29-docx-style-fix"


@app.route("/setup/<upload_id>", methods=["GET", "POST"])
def edit_setup(upload_id):
    data = _load_session_data(upload_id)
    if data is None:
        flash("Saved report progress could not be found. Please start again.", "error")
        return redirect(url_for("index"))

    meta = data["meta"]
    if request.method == "POST":
        setup = _report_setup_from_values(request.form)
        if not setup["school_name"]:
            flash("Please select at least one school.", "error")
            return render_template(
                "start.html",
                upload_id=upload_id,
                can_review=_session_can_review(meta),
                schools=_school_options(),
                aquatrack_url=os.environ.get("AQUATRACK_URL", "/"),
                **setup,
            ), 400
        setup_changed = any(meta.get(key, "") != value for key, value in setup.items())
        meta.update(setup)
        if setup_changed:
            meta["fixture_ids"] = []
            meta["unknown_fixtures"] = []
        compose_draft = dict(meta.get("compose_draft") or {})
        if compose_draft:
            compose_draft["school_name"] = setup["school_name"]
            meta["compose_draft"] = compose_draft
        if setup_changed and _session_has_style(meta):
            meta["style_needs_refresh"] = True
        _save_session_data(upload_id, data["samples"], meta)
        return redirect(url_for("reference_upload", upload_id=upload_id))

    return render_template(
        "start.html",
        upload_id=upload_id,
        campus_id=meta.get("campus_id", ""),
        campus_ids=meta.get("campus_ids") or ([meta.get("campus_id")] if meta.get("campus_id") else []),
        school_name=meta.get("school_name", ""),
        report_style=meta.get("report_style", "wa_school"),
        can_review=_session_can_review(meta),
        schools=_school_options(),
        aquatrack_url=os.environ.get("AQUATRACK_URL", "/"),
    )


@app.route("/upload-options", methods=["GET", "POST"])
def upload_options():
    setup = _report_setup_from_values(request.values)
    if request.method == "GET":
        return redirect(url_for("index"))
    if not setup["school_name"]:
        flash("Please select at least one school.", "error")
        return redirect(url_for("index"))

    upload_id = uuid.uuid4().hex[:12]
    session["selected_fixture_ids"] = []
    _save_session_data(upload_id, [], {
        **setup,
        "fixture_ids": [],
        "filename": "AquaTrack lead testing records",
        "uploaded_at": datetime.now().isoformat(),
        "unknown_fixtures": [],
        "original_files": [],
        "style_draft": {},
        "style_adaptation_error": "",
        "style_needs_refresh": False,
    })
    return redirect(url_for("reference_upload", upload_id=upload_id))


@app.route("/reference-upload", defaults={"upload_id": None}, methods=["GET", "POST"])
@app.route("/reference-upload/<upload_id>", methods=["GET", "POST"])
def reference_upload(upload_id=None):
    saved = _load_session_data(upload_id) if upload_id else None
    if not upload_id or saved is None:
        flash("Saved report progress could not be found. Please start again.", "error")
        return redirect(url_for("index"))

    meta = saved["meta"]
    setup = {
        "report_style": "wa_school",
        "organization": meta.get("organization", ""),
        "school_name": meta.get("school_name", ""),
        "campus_id": meta.get("campus_id", ""),
        "campus_ids": meta.get("campus_ids") or ([meta.get("campus_id")] if meta.get("campus_id") else []),
    }
    if not setup["school_name"]:
        flash("Please select at least one school first.", "error")
        return redirect(url_for("index"))
    fixtures = _fixtures_for_campuses(setup["campus_ids"])
    selected_ids = set(meta.get("fixture_ids") or [row["id"] for row in fixtures])
    template_context = {
        **setup,
        "upload_id": upload_id,
        "fixtures": fixtures,
        "selected_fixture_ids": selected_ids,
        "can_review": _session_can_review(meta),
    }
    if request.method == "GET":
        return render_template("reference_upload.html", **template_context)

    fixture_ids = request.form.getlist("fixture_ids")
    if not fixture_ids:
        flash("Please select at least one fixture.", "error")
        return render_template("reference_upload.html", **template_context), 400

    try:
        samples, selected_fixtures = _samples_from_inventory(
            setup["campus_ids"],
            fixture_ids,
        )
    except Exception as e:
        flash(f"Could not load the selected fixtures: {e}", "error")
        return render_template("reference_upload.html", **template_context), 400

    session["selected_fixture_ids"] = fixture_ids
    unknown = sorted(set(registry.unknown_ids([s.fixture_id for s in samples])))
    updated_meta = {
        **meta,
        "fixture_ids": fixture_ids,
        "filename": "AquaTrack lead testing records",
        "uploaded_at": datetime.now().isoformat(),
        "unknown_fixtures": unknown,
        "selected_fixture_count": len(selected_fixtures),
        "parser_source": "AquaTrack database",
        "style_draft": {},
        "style_adaptation_error": "",
        "style_needs_refresh": _session_has_style(meta),
        **setup,
    }
    _save_session_data(upload_id, samples, updated_meta)
    return redirect(url_for("report_style", upload_id=upload_id))


@app.route("/report-style/<upload_id>", methods=["GET", "POST"])
def report_style(upload_id):
    data = _load_session_data(upload_id)
    if data is None:
        flash("Testing-result upload could not be found. Please start again.", "error")
        return redirect(url_for("index"))

    samples = data["samples"]
    meta = data["meta"]
    setup = {
        "report_style": meta.get("report_style", "wa_school"),
        "organization": meta.get("organization", ""),
        "school_name": meta.get("school_name", ""),
    }
    existing_style = _original_file_for(meta, "style_report_file")
    existing_header = _original_file_for(meta, "header_template_file")
    template_context = {
        "upload_id": upload_id,
        "fixture_count": len(samples),
        "result_filename": meta.get("filename", "the testing result"),
        "existing_style": existing_style,
        "existing_header": existing_header,
        "can_review": (
            existing_style is not None
            and existing_header is not None
            and not meta.get("style_needs_refresh", False)
        ),
        **setup,
    }
    if request.method == "GET":
        return render_template("report_style.html", **template_context)

    style_file = request.files.get("style_report_file")
    if not style_file or not style_file.filename:
        if not existing_style:
            flash("Please upload the required sample of your report style.", "error")
            return render_template("report_style.html", **template_context), 400
        style_path = Path(existing_style["path"])
        if not style_path.exists():
            flash("The saved report-style sample could not be found. Please upload it again.", "error")
            return render_template("report_style.html", **template_context), 400
        style_bytes = style_path.read_bytes()
        style_filename = existing_style["filename"]
        saved_style = existing_style
    else:
        style_bytes = style_file.read()
        style_file.stream.seek(0)
        style_filename = style_file.filename
        saved_style = _save_original_upload(upload_id, "style_report_file", style_file)

    header_file = request.files.get("header_template_file")
    if not header_file or not header_file.filename:
        if not existing_header:
            flash("Please upload the required Header template.", "error")
            return render_template("report_style.html", **template_context), 400
        header_path = Path(existing_header["path"])
        if not header_path.exists():
            flash("The saved Header template could not be found. Please upload it again.", "error")
            return render_template("report_style.html", **template_context), 400
        saved_header = existing_header
    else:
        header_bytes = header_file.read()
        try:
            _validate_header_template(header_bytes, header_file.filename)
        except ValueError as exc:
            flash(str(exc), "error")
            return render_template("report_style.html", **template_context), 400
        header_file.stream.seek(0)
        saved_header = _save_original_upload(
            upload_id, "header_template_file", header_file
        )

    originals = list(meta.get("original_files", []))
    originals = [
        entry
        for entry in originals
        if entry.get("field") not in {"style_report_file", "header_template_file"}
    ]
    if saved_style:
        originals.append(saved_style)
    if saved_header:
        originals.append(saved_header)

    style_draft = {}
    style_adaptation_error = ""
    try:
        style_draft = draft_school_communication_from_reference(
            style_bytes,
            style_filename,
            _communication_facts(samples, setup),
        )
    except ClaudeStyleError as exc:
        style_adaptation_error = str(exc)

    meta.update({
        "original_files": originals,
        "style_draft": style_draft,
        "style_adaptation_error": style_adaptation_error,
        "style_needs_refresh": False,
    })
    _save_session_data(upload_id, samples, meta)
    return redirect(url_for("compose", upload_id=upload_id))


@app.route("/upload", methods=["POST"])
def upload():
    setup_id = request.form.get("setup_id", "").strip()
    staged_references = []
    if setup_id:
        staged = _load_reference_setup(setup_id)
        if not staged:
            flash("Reference upload could not be found. Please start again.", "error")
            return redirect(url_for("index"))
        setup = staged["setup"]
        staged_references = staged.get("reference_files", [])
    else:
        setup = _report_setup_from_values(request.form)
    f = request.files.get("file")
    if not f or not f.filename:
        flash("Please choose a file.", "error")
        return redirect(_current_upload_url(setup, setup_id))
    if setup["report_style"] == "wa_school" and not staged_references:
        flash("Please upload the required reference samples first.", "error")
        return redirect(url_for("reference_upload", **setup))

    parsed_bytes = f.read()
    try:
        samples = _parse_upload(f.filename, parsed_bytes)
    except Exception as e:
        flash(f"Could not parse file: {e}", "error")
        return redirect(_current_upload_url(setup, setup_id))

    if not samples:
        flash("No samples found in file.", "error")
        return redirect(_current_upload_url(setup, setup_id))

    sample_fixture_ids = [s.fixture_id for s in samples]
    unknown = sorted(set(registry.unknown_ids(sample_fixture_ids)))

    upload_id = uuid.uuid4().hex[:12]
    originals = []
    parsed_path = WORK / f"originals_{upload_id}"
    parsed_path.mkdir(exist_ok=True)
    parsed_safe = secure_filename(f.filename) or "parsed_upload"
    parsed_file_path = parsed_path / f"parsed_{parsed_safe}"
    parsed_file_path.write_bytes(parsed_bytes)
    originals.append(_source_file_entry("parsed", f.filename, parsed_file_path))
    originals.extend(staged_references)
    for field in ("coc_file",):
        saved = _save_original_upload(upload_id, field, request.files.get(field))
        if saved:
            originals.append(saved)
    _save_session_data(upload_id, samples, {
        "filename": f.filename,
        "uploaded_at": datetime.now().isoformat(),
        "unknown_fixtures": unknown,
        "original_files": originals,
        **setup,
    })
    return redirect(url_for("compose", upload_id=upload_id))


@app.route("/original/<upload_id>/<int:file_idx>")
def original_file(upload_id, file_idx: int):
    data = _load_session_data(upload_id)
    if data is None:
        abort(404)
    files = data.get("meta", {}).get("original_files", [])
    if file_idx < 0 or file_idx >= len(files):
        abort(404)
    entry = files[file_idx]
    path = Path(entry["path"])
    if not path.exists():
        abort(404)
    return send_file(path, as_attachment=False, download_name=entry["filename"])


@app.get("/original/<upload_id>/<int:file_idx>/page/<int:page_number>.png")
def original_pdf_page(upload_id, file_idx: int, page_number: int):
    """Render a PDF page as PNG so Review never relies on browser PDF plugins."""
    data = _load_session_data(upload_id)
    if data is None:
        abort(404)
    files = data.get("meta", {}).get("original_files", [])
    if file_idx < 0 or file_idx >= len(files):
        abort(404)
    entry = files[file_idx]
    path = Path(entry["path"])
    if path.suffix.lower() != ".pdf" or not path.exists():
        abort(404)

    import pypdfium2 as pdfium

    document = pdfium.PdfDocument(str(path))
    try:
        if page_number < 1 or page_number > len(document):
            abort(404)
        page = document[page_number - 1]
        try:
            bitmap = page.render(scale=1.8)
            try:
                image = bitmap.to_pil().convert("RGB").copy()
            finally:
                bitmap.close()
        finally:
            page.close()
    finally:
        document.close()

    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return send_file(
        output,
        mimetype="image/png",
        as_attachment=False,
        download_name=f"{path.stem}-page-{page_number}.png",
        max_age=300,
    )


def _compose_draft_from_form() -> dict:
    contacts = []
    for i in (1, 2):
        contacts.append({
            "name": request.form.get(f"contact{i}_name", "").strip(),
            "title": request.form.get(f"contact{i}_title", "").strip(),
            "phone": request.form.get(f"contact{i}_phone", "").strip(),
            "email": request.form.get(f"contact{i}_email", "").strip(),
        })
    selected_buildings = [
        value.strip() for value in request.form.getlist("building") if value.strip()
    ]
    return {
        "building": selected_buildings[0] if selected_buildings else "",
        "buildings": selected_buildings,
        "sampling_dates": request.form.get("sampling_dates", "").strip(),
        "organization": request.form.get("organization", "").strip(),
        "school_name": request.form.get("school_name", "").strip(),
        "intro": request.form.get("introduction", "").strip(),
        "actions": request.form.get("actions_taken", "").strip(),
        "notes": request.form.get("notes", "").strip(),
        "contacts": contacts,
    }


@app.route("/compose/<upload_id>/save-draft", methods=["POST"])
def save_compose_draft(upload_id):
    data = _load_session_data(upload_id)
    if data is None:
        return jsonify({"saved": False, "error": "Report progress was not found."}), 404
    _apply_table_edits(data["samples"])
    data["meta"]["compose_draft"] = _compose_draft_from_form()
    _save_session_data(upload_id, data["samples"], data["meta"])
    return jsonify({"saved": True})


@app.route("/compose/<upload_id>", methods=["GET", "POST"])
def compose(upload_id):
    data = _load_session_data(upload_id)
    if data is None:
        abort(404)

    samples = data["samples"]
    unknown = data["meta"]["unknown_fixtures"]
    original_files = data["meta"].get("original_files", [])
    preview_field_order = {
        "style_report_file": 0,
        "header_template_file": 1,
        "school_testing_result_file": 2,
    }
    preview_files = sorted(
        (
            {**entry, "file_idx": file_idx}
            for file_idx, entry in enumerate(original_files)
            if entry.get("field") in preview_field_order
        ),
        key=lambda entry: preview_field_order[entry.get("field")],
    )
    # AquaTrack's Communication workspace is district-only. Keep the refined
    # reporter workflow while preventing legacy report modes from resurfacing
    # through saved session metadata.
    initial_report_style = "wa_school"
    compose_draft = data["meta"].get("compose_draft") or {}
    initial_organization = compose_draft.get(
        "organization", data["meta"].get("organization", "")
    )
    initial_school_name = data["meta"].get("school_name", "")
    style_draft = data["meta"].get("style_draft") or {}
    style_adaptation_error = data["meta"].get("style_adaptation_error", "")
    for entry in preview_files:
        is_school_result = entry.get("field") == "school_testing_result_file"
        if "preview" not in entry or is_school_result:
            entry["preview"] = _build_source_preview(
                Path(entry["path"]),
                entry["filename"],
                result_pages_only=is_school_result,
                relevant_pages_override=entry.get("relevant_pages"),
            )

    if request.method == "POST":
        return _handle_compose_post(upload_id, samples)

    # GET: render the compose form

    # Detected buildings only. Do not populate the compose screen with the
    # entire UW alias list.
    detected = detected_buildings_from_samples(samples)

    buildings: list[dict] = []
    for d in detected:
        buildings.append({
            "name": d["name"],
            "code": d["code"],
            "is_detected": True,
            "is_unknown": d["source"] == "unknown",
            "sample_count": d["sample_count"],
        })
    if not buildings:
        buildings.append({
            "name": "Unknown Building",
            "code": "",
            "is_detected": False,
            "is_unknown": True,
            "sample_count": len(samples),
        })

    available_buildings = {entry["name"] for entry in buildings}
    saved_buildings = compose_draft.get("buildings") or []
    if not saved_buildings and compose_draft.get("building"):
        saved_buildings = [compose_draft["building"]]
    selected_buildings = [
        name for name in saved_buildings if name in available_buildings
    ] or [entry["name"] for entry in buildings]
    suggested = selected_buildings[0]

    # Decide which Actions Taken template to prefill. We use the same default
    # threshold profile as the preview table so the actions match what the
    # author sees highlighted.
    levels_default = load_profile(DEFAULT_PROFILE_NAME)
    has_exceedances = any(
        m.exceeds(level.threshold)
        for s in samples
        for m in s.measurements
        for level in levels_default
        if level.analyte == m.analyte
    )

    # Build defaults map keyed by building name (so the JS switcher works).
    # Passing `samples` lets _default_intro produce a per-building intro
    # naming the actual fixture count, fixture type, and analytes detected.
    default_display_name = (
        initial_school_name if initial_report_style == "wa_school" else ""
    )
    defaults_map = {}
    for entry in buildings:
        defaults = defaults_for(
            entry["name"], has_exceedances, samples,
            report_style=initial_report_style,
            display_name=default_display_name,
        )
        if style_draft:
            defaults.update({
                "intro": style_draft.get("intro", defaults["intro"]),
                "actions": style_draft.get("actions", defaults["actions"]),
                "notes": style_draft.get("notes", ""),
            })
        defaults_map[entry["name"]] = defaults

    suggested_defaults = (defaults_map.get(suggested) or {
        "intro": _default_intro(
            {}, None, default_display_name or suggested or "",
            neutral=initial_report_style == "wa_school",
        ),
        "actions": _default_actions(
            {}, has_exceedances, neutral=initial_report_style == "wa_school",
        ),
        "contacts": _default_contacts(
            {}, neutral=initial_report_style == "wa_school",
        ),
    })
    if compose_draft:
        suggested_defaults = dict(suggested_defaults)
        suggested_defaults.update({
            "intro": compose_draft.get("intro", suggested_defaults.get("intro", "")),
            "actions": compose_draft.get("actions", suggested_defaults.get("actions", "")),
            "notes": compose_draft.get("notes", suggested_defaults.get("notes", "")),
            "contacts": compose_draft.get(
                "contacts", suggested_defaults.get("contacts", [])
            ),
        })
        defaults_map[suggested] = suggested_defaults

    # Codes that need confirmation: detected but not in profile/alias
    unprofiled_codes = [d for d in detected if d["source"] == "unknown"]

    return render_template(
        "compose.html",
        upload_id=upload_id,
        samples=samples,
        buildings=buildings,
        selected_buildings=selected_buildings,
        suggested_building=suggested,
        suggested_defaults=suggested_defaults,
        defaults_map_json=json.dumps(defaults_map),
        unknown=unknown,
        inferred_code=_infer_building_code_from_samples(samples),
        detected_buildings=detected,
        is_multi_building=len(detected) > 1,
        has_exceedances=has_exceedances,
        unprofiled_codes=unprofiled_codes,
        sampling_date_range=(
            compose_draft.get("sampling_dates") or _sample_date_range(samples)
        ),
        nomenclature_help=NOMENCLATURE_HELP,
        preview_files=preview_files,
        reference_style_applied=bool(style_draft),
        reference_style_summary=style_draft.get("style_summary", ""),
        style_adaptation_error=style_adaptation_error,
        initial_report_style=initial_report_style,
        initial_organization=initial_organization,
        initial_school_name=initial_school_name,
        report_school_name=(
            compose_draft.get("school_name") or initial_school_name or suggested
        ),
    )


def _building_label(rec: dict) -> str:
    if not rec.get("code"):
        return rec.get("name", "Unknown Building")
    name = rec.get("name", "")
    code = rec["code"]
    return name if code in name else f"{name} ({code})"


def _fixture_location_label(fixture, fixture_id: str, known: bool) -> str:
    code = _building_code_for(fixture_id) or (fixture_id or "").split("_")[0].upper()
    room = _normalize_room_label(getattr(fixture, "room", "")) if known else ""
    if room:
        return f"{code} room {room}"
    if not known:
        return _unknown_fixture_location_label(fixture_id)
    return code or fixture_id


def _unknown_fixture_location_label(fixture_id: str) -> str:
    code = _building_code_for(fixture_id) or (fixture_id or "").split("_")[0].upper()
    if not code:
        return ""
    parts = (fixture_id or "").split("_")
    room = ""
    for part in reversed(parts[1:]):
        token = part.strip()
        if not token:
            continue
        upper = token.upper()
        if upper in {"PF", "MF", "BRS", "WBF"}:
            continue
        if upper in {"250ML", "1L"}:
            continue
        if any(ch.isdigit() for ch in upper):
            room = upper
            break
    return f"{code} room {room}" if room else code


def _normalize_room_label(room: str) -> str:
    text = (room or "").strip()
    for prefix in ("Rm ", "Room ", "rm ", "room "):
        if text.startswith(prefix):
            text = text[len(prefix):]
            break
    return text.strip()


def _slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")


def _ordered_analytes_for(samples, *, include_missing_defaults: bool) -> list[str]:
    analytes: list[str] = []
    if include_missing_defaults:
        analytes.extend(REPORT_ANALYTES)
    for idx, s in enumerate(samples):
        for m in s.measurements:
            if m.analyte not in analytes:
                analytes.append(m.analyte)
    return analytes


def _preview_rows(samples) -> list[dict]:
    """Build the report-like highlighted preview table."""
    rows = []
    levels_default = load_profile(DEFAULT_PROFILE_NAME)

    def preview_display(m) -> str:
        if m is None:
            return "not reported"
        if m.below_dl and m.detection_limit is None:
            return "not detected"
        return m.display

    for idx, s in enumerate(samples):
        code = _building_code_for(s.fixture_id)
        reported_building = _reported_building_name(s)
        building = {
            "name": reported_building or "Unknown Building",
            "code": "",
        } if reported_building else (
            resolve_building(code) if code else {"name": "Unknown Building", "code": ""}
        )
        fixture = registry.get(s.fixture_id)
        known = fixture is not None
        if fixture is None:
            fixture = _placeholder_fixture(s.fixture_id)
        sev = evaluate_sample(s, levels_default)
        cells = {}
        for analyte in _ordered_analytes_for(samples, include_missing_defaults=True):
            m = s.measurement(analyte)
            cells[analyte] = {
                "display": preview_display(m),
                "severity": sev.get(analyte, "ok") if m else "ok",
                "input_name": f"result_{idx}_{_slug(analyte)}",
            }
        rows.append({
            "idx": idx,
            "sample_id": s.client_sample_id,
            "building": _building_label(building),
            "sample_volume": _format_sample_volume(s.volume_ml),
            "fixture_location": getattr(s, "fixture_label", "") or _fixture_location_label(
                fixture, s.fixture_id, known),
            "cells": cells,
        })
    return rows


def _parse_volume_label(label: str) -> int:
    s = (label or "").strip().lower()
    if not s or s == "not specified":
        return 0
    m = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*(ml|l)\b", s)
    if not m:
        return 0
    value = float(m.group(1))
    return int(round(value * 1000)) if m.group(2) == "l" else int(round(value))


def _apply_table_edits(samples) -> None:
    """Apply editable preview-table fields before DOCX generation."""
    analytes = _ordered_analytes_for(samples, include_missing_defaults=True)
    for idx, sample in enumerate(samples):
        sample_id = request.form.get(f"sample_{idx}_id", sample.client_sample_id).strip()
        if sample_id:
            sample.client_sample_id = sample_id
            sample.fixture_id, parsed_volume = _parse_client_id(sample_id)
            if parsed_volume:
                sample.volume_ml = parsed_volume
        volume_label = request.form.get(f"sample_{idx}_volume", "").strip()
        if volume_label:
            sample.volume_ml = _parse_volume_label(volume_label)
        fixture_label = request.form.get(f"sample_{idx}_fixture_label", "").strip()
        if fixture_label:
            sample.fixture_label = fixture_label

        existing = {m.analyte: m for m in sample.measurements}
        edited: list[Measurement] = []
        for analyte in analytes:
            field = f"result_{idx}_{_slug(analyte)}"
            if field not in request.form:
                m = existing.get(analyte)
                if m:
                    edited.append(m)
                continue
            raw = request.form.get(field, "").strip()
            if raw.lower() in {"", "not detected", "nd", "n.d.", "not reported"}:
                continue
            unit = existing.get(analyte).unit if existing.get(analyte) else None
            try:
                edited.append(_parse_value(raw, analyte, unit))
            except ValueError:
                m = existing.get(analyte)
                if m:
                    edited.append(m)
        sample.measurements = edited


def _format_date(d: date) -> str:
    return f"{d.strftime('%B')} {d.day}, {d.year}"


def _sample_date_range(samples) -> str:
    dates = sorted({
        d
        for s in samples
        for d in (s.collection_date, s.analysis_date)
        if d is not None
    })
    if not dates:
        return ""
    if dates[0] == dates[-1]:
        return _format_date(dates[0])
    return f"{_format_date(dates[0])} through {_format_date(dates[-1])}"


def _docx_analytes_for(samples) -> list[str]:
    """School Review and exported DOCX display Lead results only."""
    return ["Lead"] if any(s.measurement("Lead") is not None for s in samples) else []


def _handle_compose_post(upload_id: str, samples):
    """Process the compose form submission and return the DOCX file."""
    _apply_table_edits(samples)
    selected_buildings = [
        value.strip() for value in request.form.getlist("building") if value.strip()
    ]
    if not selected_buildings:
        detected = detected_buildings_from_samples(samples)
        selected_buildings = [entry["name"] for entry in detected]
    if not selected_buildings:
        flash("Please choose at least one building.", "error")
        return redirect(url_for("compose", upload_id=upload_id))

    # Filter samples to this building. We accept samples in two ways:
    #   1. Sample's fixture_id is registered to this building in fixtures.json
    #   2. Sample's fixture_id prefix resolves to this building name
    # The second path lets us generate draft reports for buildings whose
    # fixtures haven't been registered yet — placeholder fixtures get
    # synthesized in _build_rows.
    bldg_fixture_ids = {
        fixture.fixture_id
        for building_name in selected_buildings
        for fixture in registry.by_building(building_name)
    }
    detected = detected_buildings_from_samples(samples)

    def sample_matches_building(s) -> bool:
        if not detected and "Unknown Building" in selected_buildings:
            return True
        if _reported_building_name(s) in selected_buildings:
            return True
        if s.fixture_id in bldg_fixture_ids:
            return True
        code = _building_code_for(s.fixture_id)
        if not code:
            return False
        rec = resolve_building(code)
        return rec["name"] in selected_buildings

    filtered = [s for s in samples if sample_matches_building(s)]
    if not filtered:
        flash(
            f"No samples in this upload match {selected_buildings!r}. "
            f"Sample fixtures: {sorted({s.fixture_id for s in samples})}. "
            f"Pick a different building.",
            "error",
        )
        return redirect(url_for("compose", upload_id=upload_id))

    levels = load_profile(DEFAULT_PROFILE_NAME)
    saved = _load_session_data(upload_id) or {}
    meta = saved.get("meta", {})
    meta["compose_draft"] = _compose_draft_from_form()
    _save_session_data(upload_id, samples, meta)
    report_building = (
        request.form.get("school_name", "").strip()
        or ", ".join(selected_buildings)
    )
    report_style = "wa_school"
    organization = request.form.get("organization", "").strip()

    contacts = []
    for i in (1, 2):
        contact = {
            "name": request.form.get(f"contact{i}_name", "").strip(),
            "title": request.form.get(f"contact{i}_title", "").strip(),
            "phone": request.form.get(f"contact{i}_phone", "").strip(),
            "email": request.form.get(f"contact{i}_email", "").strip(),
        }
        if any(contact.values()):
            contacts.append(contact)

    header_entry = _original_file_for(meta, "header_template_file")
    header_template_path = Path(header_entry["path"]) if header_entry else None
    if not header_template_path or not header_template_path.exists():
        flash("The required Header template could not be found. Please upload it again.", "error")
        return redirect(url_for("report_style", upload_id=upload_id))

    ctx = ReportContext(
        building=report_building,
        report_date=None,  # blank in DOCX; author types final date in Word
        sampling_date_range=(
            request.form.get("sampling_dates", "").strip()
            or _sample_date_range(filtered)
        ),
        introduction_md=request.form.get("introduction", "").strip(),
        actions_taken_md=request.form.get("actions_taken", "").strip(),
        contacts=contacts,
        samples=filtered,
        action_levels=levels,
        analytes_shown=_docx_analytes_for(filtered),
        notes_md=request.form.get("notes", "").strip(),
        report_style=report_style,
        organization=organization,
        reference_style_applied=bool(meta.get("style_draft")),
        reference_layout=(meta.get("style_draft") or {}).get("layout", "report"),
        header_template_path=str(header_template_path),
    )

    buf = BytesIO()
    render_docx(ctx, registry, buf)
    safe_building = re.sub(r"[^A-Za-z0-9]+", "_", report_building).strip("_")
    date_part = ctx.report_date.isoformat() if ctx.report_date else "undated"
    filename = f"{safe_building}_water_results_{date_part}.docx"
    if supabase.configured:
        report_bytes = buf.getvalue()
        storage_path = f"{g.current_user['id']}/reports/{uuid.uuid4().hex}_{filename}"
        supabase.upload_bytes(
            storage_path,
            report_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        campus_ids = meta.get("campus_ids") or (
            [meta.get("campus_id")] if meta.get("campus_id") else []
        )
        supabase.insert("communication_generated_reports", {
            "created_by": g.current_user["id"],
            "campus_id": campus_ids[0] if len(campus_ids) == 1 else None,
            "fixture_ids": meta.get("fixture_ids") or [],
            "school_name": report_building,
            "file_name": filename,
            "storage_path": storage_path,
            "source_upload_id": upload_id,
        })
        buf.seek(0)
    return send_file(
        buf, as_attachment=True, download_name=filename,
        mimetype=("application/vnd.openxmlformats-"
                  "officedocument.wordprocessingml.document"),
    )


if __name__ == "__main__":
    app.run(debug=True, port=5000)
