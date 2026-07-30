"""Render a ReportContext to an editable .docx file.

DOCX is the right output format for these memos: EH&S authors will
nearly always tweak wording before sending to building occupants.
"""
from __future__ import annotations
from copy import deepcopy
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import ReportContext
from .fixtures import FixtureRegistry
from .report import _build_rows, _violation_prose


# UW purple-ish; tweak to match your brand spec
UW_PURPLE = RGBColor(0x4B, 0x2E, 0x83)
WARN_FILL = "FDE7D3"      # light orange — aesthetic exceedance
VIOLATION_FILL = "F8C7C7"  # light red — health-based exceedance
HEADER_FILL = "ECE6F5"     # light purple — table header


def _display_unit(unit: str) -> str:
    """Map canonical storage units to UW report display labels.

    UW reports use 'ppm' for non-lead metals (1 mg/L = 1 ppm in water).
    The internal Measurement.unit stays 'mg/L'; this only changes labels.
    """
    return "ppm" if unit == "mg/L" else unit


def _shade_cell(cell, hex_color: str):
    """Set background fill of a table cell. python-docx has no high-level API
    for this, so we manipulate the underlying XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Set cell padding in twips (1/20 of a point). Defaults give tight padding
    suitable for a data table — Word's default is much looser."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove any existing tcMar
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left),
                       ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tc_pr.append(tcMar)


def _set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def _set_repeat_table_header(row):
    """Repeat this row as the table header on continued pages."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _set_table_fixed_width(table, width_in: float):
    """Give Word explicit table geometry so it does not stretch columns."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(width_in * 1440)))


def _set_cell_vcenter(cell):
    """Vertically center cell content."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:vAlign"))
    if existing is not None:
        tc_pr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tc_pr.append(vAlign)


def _set_cell_text(cell, text: str, *, bold: bool = False,
                   color: RGBColor | None = None, size: Pt | None = None,
                   align: int | None = None):
    """Replace cell content with a single run, formatted as specified.

    Using cell.text leaves Normal-style runs and won't let us set color/bold
    in one call; this helper is the equivalent."""
    # Wipe existing paragraphs
    cell.text = ""  # leaves one empty paragraph
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    # Multi-line content: split on \n into multiple paragraphs
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            target = p
        else:
            target = cell.add_paragraph()
            target.paragraph_format.space_before = Pt(0)
            target.paragraph_format.space_after = Pt(0)
            target.paragraph_format.line_spacing = 1
            if align is not None:
                target.alignment = align
        run = target.add_run(line)
        if bold:
            run.bold = True
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = size


def _building_code_from_fixture_id(fixture_id: str) -> str:
    m = re.match(r"^([A-Za-z]+\d*)", fixture_id or "")
    return m.group(1).upper() if m else ""


def _normalize_room_text(room: str) -> str:
    text = (room or "").strip()
    for prefix in ("Rm ", "Room ", "rm ", "room "):
        if text.startswith(prefix):
            text = text[len(prefix):]
            break
    text = text.strip("() ")
    return text.upper() if text else ""


def _room_from_fixture_id(fixture_id: str) -> str:
    parts = (fixture_id or "").split("_")
    for token in reversed(parts[1:]):
        cleaned = token.strip()
        upper = cleaned.upper()
        if upper in {"PF", "MF", "BRS", "WBF", "250ML", "1L"}:
            continue
        if any(ch.isdigit() for ch in cleaned):
            return cleaned.upper()
    return ""


def _split_fixture_label(label: str) -> tuple[str, str]:
    text = re.sub(r"\s+", " ", label or "").strip()
    m = re.match(r"^([A-Za-z]+\d*)\s+room\s+(.+)$", text, flags=re.I)
    if m:
        return m.group(1).upper(), _normalize_room_text(m.group(2))
    m = re.match(r"^([A-Za-z]+\d*)$", text)
    if m:
        return m.group(1).upper(), ""
    return "", ""


def _table_fixture_label(row: dict) -> str:
    """Compact display for the results table location column.

    The report table should show the building/facility code on one centered
    line and the room underneath, e.g.:
      COM
      room B002
    """
    sample = row["sample"]
    fixture = row["fixture"]
    source_label = re.sub(
        r"\s+", " ", getattr(sample, "fixture_label", "") or ""
    ).strip()
    label_code, label_room = _split_fixture_label(source_label)

    # Lab/agency reports commonly provide a complete free-form description,
    # e.g. "Water Fountain - Rm27 - Fountain/Tap Combo". Preserve that source
    # text verbatim instead of discarding it merely because it does not match
    # the older "COM room B002" shorthand used by UW fixture IDs.
    if source_label and not label_code and not label_room:
        return source_label

    code = label_code or _building_code_from_fixture_id(sample.fixture_id) or fixture.building
    room = (
        label_room
        or _normalize_room_text(getattr(fixture, "room", ""))
        or _room_from_fixture_id(sample.fixture_id)
    )
    if room:
        return f"{code}\nroom {room}"
    return code or sample.fixture_id


def _table_building_label(row: dict) -> str:
    """Prefer the report's row-level Building Name for the DOCX table."""
    sample = row["sample"]
    source_fields = getattr(sample, "source_fields", {}) or {}
    for value in (
        getattr(sample, "building_name", ""),
        source_fields.get("building_name", ""),
        source_fields.get("Building Name", ""),
        source_fields.get("building", ""),
        getattr(row["fixture"], "building", ""),
    ):
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if text:
            return text
    return "not specified"


def _render_results_table(doc: Document, rows: list[dict],
                          analytes: list[str]):
    """Render the Table 1 results grid matching the UW report style.

    Layout target:
      - Simple sample number
      - Source-reported building name
      - Fixture/location label in the wider text column
      - One column per analyte, each header is "Analyte" / "(unit)" stacked
      - Tight cell padding, vertically centered, compact page flow
      - 'Table Grid' style for thin borders (no colored bands)
    """
    def unit_for(analyte: str) -> str:
        if analyte == "Lead":
            return "ppb"
        for row in rows:
            m = row["sample"].measurement(analyte)
            if m:
                return _display_unit(m.unit)
        return "ppm"

    n_cols = 3 + len(analytes)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Compact Word-friendly geometry. When only a few analytes are present,
    # keep the table narrower instead of stretching empty space across the page.
    n_analytes = len(analytes)
    if n_analytes <= 3:
        sample_w = 0.42
        building_w = 0.95
        location_w = 2.15
        analyte_w = 0.72
        header_pt = 7.5
    elif n_analytes == 4:
        sample_w = 0.42
        building_w = 0.95
        location_w = 2.05
        analyte_w = 0.68
        header_pt = 7.5
    else:  # 5 analytes — tightest case
        sample_w = 0.42
        building_w = 0.9
        location_w = 1.9
        analyte_w = 0.64
        header_pt = 7.2
    col_widths_in = [sample_w, building_w, location_w] + [analyte_w] * n_analytes
    _set_table_fixed_width(table, sum(col_widths_in))
    for i, w in enumerate(col_widths_in):
        table.columns[i].width = Inches(w)

    # ---- Header row ----
    hdr = table.rows[0]
    hdr.height = Pt(18)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_row_cant_split(hdr)
    _set_repeat_table_header(hdr)
    cells = hdr.cells
    for i, w in enumerate(col_widths_in):
        cells[i].width = Inches(w)
    _set_cell_text(cells[0], "#", bold=True, color=UW_PURPLE,
                   size=Pt(header_pt),
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(cells[1], "Building",
                   bold=True, color=UW_PURPLE,
                   size=Pt(header_pt),
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(cells[2], "Fixture / Location",
                   bold=True, color=UW_PURPLE,
                   size=Pt(header_pt),
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, analyte in enumerate(analytes):
        unit = unit_for(analyte)
        header = f"{analyte}\n({unit})" if unit else analyte
        _set_cell_text(cells[3 + i], header,
                       bold=True, color=UW_PURPLE,
                       size=Pt(header_pt),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in cells:
        _shade_cell(c, HEADER_FILL)
        _set_cell_vcenter(c)
        _set_cell_margins(c, top=35, bottom=35, left=70, right=70)

    # ---- Data rows ----
    for i, row in enumerate(rows, start=1):
        tr = table.add_row()
        tr.height = Pt(16)
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        rcells = tr.cells
        for k, w in enumerate(col_widths_in):
            rcells[k].width = Inches(w)

        body = _table_fixture_label(row)

        _set_cell_text(rcells[0], str(i),
                       size=Pt(7.8),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(rcells[1], _table_building_label(row),
                       size=Pt(7.8),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(rcells[2], body,
                       size=Pt(7.8),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for j, analyte in enumerate(analytes):
            cell_data = row["cells"][analyte]
            _set_cell_text(rcells[3 + j], cell_data["display"],
                           size=Pt(7.8),
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            sev = cell_data["severity"]
            if sev == "violation":
                _shade_cell(rcells[3 + j], VIOLATION_FILL)
            elif sev == "warn":
                _shade_cell(rcells[3 + j], WARN_FILL)

        for c in rcells:
            _set_cell_vcenter(c)
            _set_cell_margins(c, top=30, bottom=30, left=70, right=70)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = UW_PURPLE


def _ensure_renderer_styles(doc: Document):
    """Restore the built-in styles required by the report renderer.

    Header-only DOCX files exported by tools such as Pages or document
    generators can omit Word's Normal, Table Grid, and List Bullet styles.
    Copy only the missing definitions from python-docx's standard template so
    the uploaded header and media stay intact while the generated body remains
    a conventional editable Word document.
    """
    required = ("Normal", "Normal Table", "Table Grid", "List Bullet")
    existing = {style.name for style in doc.styles}
    if any(name not in existing for name in required):
        standard = Document()
        for name in required:
            if name not in existing:
                doc.styles._element.append(
                    deepcopy(standard.styles[name]._element)
                )
    return doc.styles["Normal"]


def _contact_clause(c: dict) -> str:
    """Format one contact as 'Name, Title at Phone or Email'.

    Phone/title/email are optional; only present pieces are included.
    """
    parts: list[str] = []
    if c.get("name"):
        parts.append(c["name"])
    if c.get("title"):
        parts.append(c["title"])
    head = ", ".join(parts)
    tail_bits: list[str] = []
    if c.get("phone"):
        tail_bits.append(c["phone"])
    if c.get("email"):
        tail_bits.append(c["email"])
    tail = " or ".join(tail_bits)
    if not tail:
        return head
    return f"{head} at {tail}"


def _questions_prose(ehs: dict, fac: dict) -> str:
    """Build the standard 'Questions' paragraph from two contacts.

    Matches the recurring UW EH&S memo wording: first sentence about water
    testing results, second sentence about water system / fountain replacement.
    """
    return (
        "If you have any questions about the water testing results or this "
        f"communication, contact {_contact_clause(ehs)}. "
        "For questions about the water system or fountain replacement, "
        f"contact {_contact_clause(fac)}."
    )


def _add_md_paragraphs(doc: Document, md_text: str):
    """Tiny markdown subset: paragraphs, bullets, **bold**, and *italic*."""
    if not md_text:
        return

    def add_inline(p, text: str):
        tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**"):
                run = p.add_run(tok[2:-2])
                run.bold = True
            elif tok.startswith("*") and tok.endswith("*"):
                run = p.add_run(tok[1:-1])
                run.italic = True
            else:
                p.add_run(tok)

    for para in re.split(r"\n\s*\n", md_text.strip()):
        lines = [line.strip() for line in para.splitlines() if line.strip()]
        if lines and all(re.match(r"^[-•]\s+", line) for line in lines):
            for line in lines:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, re.sub(r"^[-•]\s+", "", line))
            continue
        p = doc.add_paragraph()
        add_inline(p, para)


def _set_table_borderless(table):
    """Remove all borders from a table (for header layout tables)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Drop any existing borders block
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)


def _add_paragraph_bottom_border(paragraph, color_hex: str = "4B2E83",
                                  size_eighths: int = 12):
    """Give a paragraph a bottom border — Word's idiomatic horizontal rule.

    size_eighths is in 1/8 of a point (12 = 1.5pt).
    """
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing pBdr
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_header_with_logo(doc: Document):
    """Render the UW EH&S branded header: stylized W + title block + rule line.

    Uses a 2-column borderless table so the logo and titles sit on the same
    visual row, which Word's flow layout otherwise can't do cleanly.

    The "W" is a styled glyph approximating the UW W logo. To use the
    official UW W image instead, replace the run.add_text("W") line with
    cell._element-level image insertion (or, simpler, manually drop the
    image into the generated docx in Word).
    """
    from docx.enum.table import WD_ALIGN_VERTICAL  # local import — rarely used elsewhere

    table = doc.add_table(rows=1, cols=2)
    _set_table_borderless(table)
    table.autofit = False

    # Column widths: narrow logo column, wide title column
    table.columns[0].width = Inches(0.85)
    table.columns[1].width = Inches(6.0)
    cells = table.rows[0].cells
    cells[0].width = Inches(0.85)
    cells[1].width = Inches(6.0)
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Left cell: stylized purple "W". Replace with the UW PNG by hand
    # in Word if you want the official mark.
    p_logo = cells[0].paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_logo = p_logo.add_run("W")
    r_logo.bold = True
    r_logo.font.size = Pt(36)
    r_logo.font.color.rgb = UW_PURPLE
    # Tighten the trailing space the glyph leaves
    r_logo.font.name = "Arial Black"

    # Right cell: Title + subtitle stacked
    # python-docx leaves an empty paragraph in fresh cells; reuse it.
    p_title = cells[1].paragraphs[0]
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run("ENVIRONMENTAL HEALTH & SAFETY")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = UW_PURPLE

    p_sub = cells[1].add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    r_sub = p_sub.add_run("UNIVERSITY of WASHINGTON")
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = UW_PURPLE
    # Letter-spacing to match the wide-tracked PDF style
    rPr = r_sub._r.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "40")  # 40 twips ≈ 2pt extra spacing
    rPr.append(spacing)

    # Horizontal rule under the header — empty paragraph with bottom border.
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    _add_paragraph_bottom_border(rule, color_hex="4B2E83", size_eighths=12)


def _add_neutral_header(doc: Document, organization: str = ""):
    p = doc.add_paragraph()
    r = p.add_run(organization or "Water Quality Report")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = UW_PURPLE
    sub = doc.add_paragraph()
    sub.add_run("Lead in School Drinking Water Results").bold = True
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    _add_paragraph_bottom_border(rule, color_hex="4B2E83", size_eighths=8)


def _build_testing_info_intro(ctx: "ReportContext") -> str:
    """Paragraph 1 of Testing Information.

    Explains first-draw sampling, the highlighting convention (red = health-
    based action level, orange = secondary/aesthetic standard), and includes
    the sampling date when available.
    """
    if ctx.sampling_date_range:
        opening = (
            f"The results in Table 1 below are from samples collected on "
            f"{ctx.sampling_date_range}. The concentrations shown are for "
            f"first draw water samples, meaning that the water samples were "
            f"collected after ensuring the water was left stagnant/unused "
            f"in the water fixture."
        )
    else:
        opening = (
            "The results summarized in Table 1 below were provided to EH&S "
            "after laboratory analysis. The concentrations shown are for "
            "first draw water samples, meaning that the water samples were "
            "collected after ensuring the water was left stagnant/unused "
            "in the water fixture."
        )
    return (
        opening + " "
        "Highlighted samples in Table 1 show drinking water outlets that exceed "
        "state or federal regulatory limits for drinking water. Red "
        "highlights indicate results above the default health-based action "
        "level. Orange highlights indicate results above a "
        "secondary/aesthetic standard, such as taste, odor, staining, or "
        "discoloration."
    )


_LEAD_REGULATORY_CONTEXT = (
    "The current EPA and WA State action level for lead in public drinking "
    "water is 15 parts per billion (ppb). A new EPA rule, the Lead and "
    "Copper Rule Improvements (LCRI), was finalized in 2024 to lower the "
    "lead action level in drinking water to 10 ppb for water systems to "
    "implement by 2027. WA State has additional drinking water requirements "
    "for K-12 schools, including a lower lead action level of 5 ppb, due "
    "to the higher risk that lead presents for developing children."
)


def _count_phrase(n: int, singular: str, plural: str | None = None) -> str:
    """'1 fixture' / '2 fixtures' / 'multiple fixtures' for n>=4."""
    plural = plural or (singular + "s")
    if n == 1:
        return f"one {singular}"
    if n <= 3:
        words = {2: "two", 3: "three"}
        return f"{words[n]} {plural}"
    return f"multiple {plural}"


def _highest_lead_fixture(samples, registry, threshold: float):
    """Return (fixture, max_lead_value) for the fixture with the highest lead
    reading among samples that exceed `threshold`. Returns (None, 0) if none.

    With 250mL/1L paired samples per fixture we collapse to one fixture
    record by taking the max value seen across that fixture's samples.
    """
    by_fixture: dict[str, float] = {}
    for s in samples:
        m = s.measurement("Lead")
        if not m or m.below_dl or m.value is None:
            continue
        if m.value <= threshold:
            continue
        prev = by_fixture.get(s.fixture_id, 0.0)
        if m.value > prev:
            by_fixture[s.fixture_id] = m.value
    if not by_fixture:
        return None, 0.0
    fid = max(by_fixture, key=lambda k: by_fixture[k])
    fixture = registry.get(fid)
    if fixture is None:
        # fall back to a placeholder so the prose still has a name
        from .report import _placeholder_fixture
        fixture = _placeholder_fixture(fid)
    return fixture, by_fixture[fid]


def _location_phrase(fixture, sample_id: str = "") -> str:
    """Render 'in <fixture_type> <location>' for the findings prose.

    Examples:
      'in one porcelain fountain in the basement, near room 25'
      'in one metal fountain on the second floor, near room 213'
    Falls back to the sample ID if location info is empty.
    """
    if fixture is None:
        return f"in fixture {sample_id}" if sample_id else ""

    # Floor: "Basement" -> "in the basement"; "First Floor" -> "on the first floor"
    floor = (fixture.floor or "").strip()
    floor_phrase = ""
    if floor:
        if floor.lower() == "basement":
            floor_phrase = "in the basement"
        elif floor.lower() == "mezzanine":
            floor_phrase = "on the mezzanine"
        elif "floor" in floor.lower():
            floor_phrase = "on the " + floor.lower()
        else:
            # e.g. arbitrary string the registry produced
            floor_phrase = f"on {floor}"

    # Room: "Rm 25" -> "near room 25"
    room = (fixture.room or "").strip()
    room_phrase = ""
    if room:
        # Strip leading "Rm " if present, normalize to "near room X"
        room_num = room
        for prefix in ("Rm ", "Room ", "rm ", "room "):
            if room_num.startswith(prefix):
                room_num = room_num[len(prefix):]
                break
        # Drop leading zeros: "025" -> "25"
        if room_num.isdigit():
            room_num = str(int(room_num))
        room_phrase = f"near room {room_num}"

    bits = [f"in one {fixture.fixture_type}"]
    if floor_phrase and room_phrase:
        bits.append(f"{floor_phrase}, {room_phrase}")
    elif floor_phrase:
        bits.append(floor_phrase)
    elif room_phrase:
        bits.append(room_phrase)
    return " ".join(bits)


def _action_level_label(threshold: float, source: str) -> str:
    """Format the regulatory phrase for the lead-exceedance sentence.

    Adapts based on the source string so the same code reads correctly
    whether the profile is WA-state-driven, EPA-LCRI-driven, or LCR-driven.

    Order matters: LCRI must be checked before LCR because the LCRI source
    string contains 'Lead and Copper Rule Improvements' which would
    otherwise match the LCR branch first.
    """
    src = (source or "").upper()
    if "RCW" in src:
        return f"the WA state action level of {threshold:g} ppb"
    if "IMPROVEMENTS" in src or "LCRI" in src:
        return f"the EPA LCRI action level of {threshold:g} ppb"
    if "LEAD AND COPPER" in src or "LCR" in src:
        return f"the EPA Lead and Copper Rule action level of {threshold:g} ppb"
    return f"the action level of {threshold:g} ppb"


def _build_findings_summary(ctx: "ReportContext", registry) -> str:
    """Paragraph that goes after the table footnote.

    Reports which metals exceeded thresholds, in UW prose style. Always
    mentions copper/manganese/zinc when they're below limits. Calls out
    iron aesthetic exceedances and lead health-based exceedances with
    counts and (for single-fixture lead violations) location.

    Lead handling: counts exceedances against the LOWEST lead threshold in
    the profile, regardless of severity bucket. This is necessary because
    tiered profiles assign severity 'warn' to the 5 ppb threshold and
    'violation' to the 15 ppb threshold — but the report should still call
    out Pb > 5 as a finding worth describing.
    """
    from .thresholds import evaluate_measurement

    # Per-analyte exceedance counts (unique fixtures, not samples).
    # Used for non-Pb analytes where the existing severity-based logic is fine.
    def count_exceeding(analyte: str, severity_filter: set[str] | None = None) -> int:
        seen: set[str] = set()
        for s in ctx.samples:
            m = s.measurement(analyte)
            if not m:
                continue
            sev = evaluate_measurement(m, ctx.action_levels)
            if severity_filter and sev not in severity_filter:
                continue
            if sev != "ok":
                seen.add(s.fixture_id)
        return len(seen)

    def find_threshold(analyte: str, severity: str) -> float | None:
        for level in ctx.action_levels:
            if level.analyte == analyte and level.severity == severity:
                return level.threshold
        return None

    # ---- Lead: count against the lowest threshold in the profile ----
    pb_levels = sorted(
        [l for l in ctx.action_levels if l.analyte == "Lead"],
        key=lambda l: l.threshold,
    )
    pb_threshold: float | None = None
    pb_source: str = ""
    pb_count = 0
    if pb_levels:
        lowest = pb_levels[0]
        pb_threshold = lowest.threshold
        pb_source = lowest.source
        seen: set[str] = set()
        for s in ctx.samples:
            m = s.measurement("Lead")
            if not m or m.below_dl or m.value is None:
                continue
            if m.value > pb_threshold and s.fixture_id not in seen:
                seen.add(s.fixture_id)
                pb_count += 1

    lead_only = set(ctx.analytes_shown) == {"Lead"}
    fe_warn_count = 0 if lead_only else count_exceeding("Iron", {"warn"})
    cu_count = 0 if lead_only else count_exceeding("Copper")
    mn_count = 0 if lead_only else count_exceeding("Manganese")
    zn_count = 0 if lead_only else count_exceeding("Zinc")
    fe_threshold = find_threshold("Iron", "warn")

    # If absolutely nothing was elevated, single calm sentence.
    if (pb_count == 0 and fe_warn_count == 0
            and cu_count == 0 and mn_count == 0 and zn_count == 0):
        if lead_only:
            return "Lead levels measured were below the review levels used in this report."
        return ("Lead, iron, copper, manganese, and zinc levels measured "
                "were below the default review levels and secondary "
                "standards used in this report.")

    sentences: list[str] = []

    # Cu/Mn/Zn below limits — only mention if all three are clean
    if not lead_only and cu_count == 0 and mn_count == 0 and zn_count == 0:
        sentences.append(
            "Copper, manganese, and zinc levels measured were all below "
            "federal regulatory limits."
        )

    # Iron aesthetic exceedance
    if fe_warn_count > 0 and fe_threshold is not None:
        # Use the dominant fixture type from the violators for natural prose
        iron_fixtures = []
        for s in ctx.samples:
            m = s.measurement("Iron")
            if m and not m.below_dl and m.value is not None and m.value > fe_threshold:
                iron_fixtures.append(s.fixture_id)
        type_words = []
        for fid in set(iron_fixtures):
            f = registry.get(fid)
            if f and f.fixture_type:
                type_words.append(f.fixture_type)
        from collections import Counter
        if type_words:
            dominant_type = Counter(type_words).most_common(1)[0][0]
            phrase = _count_phrase(fe_warn_count, dominant_type)
        else:
            phrase = _count_phrase(fe_warn_count, "fixture")
        sentences.append(
            f"Iron was elevated above {fe_threshold:g} ppm, the secondary "
            f"maximum contaminant limit, in {phrase}, which may contribute "
            f"to metallic taste, odor, staining, or discoloration."
        )

    # Lead exceedance — fires whenever Pb > the lowest profile threshold,
    # regardless of whether that threshold's severity is 'warn' or 'violation'.
    if pb_count > 0 and pb_threshold is not None:
        label = _action_level_label(pb_threshold, pb_source)
        if pb_count == 1:
            fixture, _ = _highest_lead_fixture(ctx.samples, registry,
                                                pb_threshold)
            loc = _location_phrase(fixture)
            if loc:
                sentences.append(f"Lead was above {label} {loc}.")
            else:
                sentences.append(f"Lead was above {label} in one fixture.")
        else:
            phrase = _count_phrase(pb_count, "fixture")
            sentences.append(f"Lead was above {label} in {phrase}.")

    pb_immediate_seen: set[str] = set()
    for s in ctx.samples:
        m = s.measurement("Lead")
        if m and not m.below_dl and m.value is not None and m.value > 15.0:
            pb_immediate_seen.add(s.fixture_id)
    if pb_immediate_seen:
        phrase = _count_phrase(len(pb_immediate_seen), "fixture")
        sentences.append(
            f"Lead was above 15 ppb in {phrase}; these results should be "
            "treated as an immediate shutoff concern while follow-up actions "
            "are confirmed."
        )

    # Closing sentence if anything elevated
    if pb_count > 0 or fe_warn_count > 0:
        sentences.append(
            "These results suggest that lead may be leaching from plumbing "
            "materials and/or water fixtures."
            if lead_only else
            "These results suggest that lead and/or other metals may be "
            "leaching from plumbing materials and/or water fixtures."
        )

    return " ".join(sentences)


def _document_from_pdf_header(path: str | Path) -> Document:
    """Create a Word document using the first PDF page as letterhead.

    Standard portrait pages are cropped to their top 2.5 inches so a full-page
    letterhead PDF does not cover the report body. Short or landscape PDFs are
    treated as header-only artwork and used in full.
    """
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(str(path))
    try:
        if len(pdf) < 1:
            raise ValueError("The PDF header template does not contain a page.")
        page = pdf[0]
        try:
            page_width_pt, page_height_pt = page.get_size()
            bitmap = page.render(scale=2.5)
            try:
                image = bitmap.to_pil().convert("RGB").copy()
            finally:
                bitmap.close()
        finally:
            page.close()
    finally:
        pdf.close()

    if page_height_pt / page_width_pt > 0.9:
        header_height_pt = min(page_height_pt, 180.0)
        crop_height = max(
            1,
            round(image.height * header_height_pt / page_height_pt),
        )
        image = image.crop((0, 0, image.width, crop_height))

    image_stream = BytesIO()
    image.save(image_stream, format="PNG", optimize=True)
    image_stream.seek(0)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.12)

    available_width_inches = (
        section.page_width - section.left_margin - section.right_margin
    ) / 914400
    rendered_height_inches = (
        available_width_inches * image.height / image.width
    )
    section.top_margin = Inches(max(1.0, rendered_height_inches + 0.3))

    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(
        image_stream,
        width=Inches(available_width_inches),
    )
    return doc


def render_docx(ctx: ReportContext, registry: FixtureRegistry,
                output: str | Path | BytesIO) -> Path | BytesIO:
    """Render the report. Pass a path to write to disk, or a BytesIO for
    in-memory (e.g. Flask send_file)."""
    header_template_path = getattr(ctx, "header_template_path", "")
    if header_template_path:
        if Path(header_template_path).suffix.lower() == ".pdf":
            doc = _document_from_pdf_header(header_template_path)
        else:
            doc = Document(str(header_template_path))
        body = doc._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
    else:
        doc = Document()
    section = doc.sections[0]
    if not header_template_path:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Tighten default style a bit
    style = _ensure_renderer_styles(doc)
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    reference_letter = (
        getattr(ctx, "reference_style_applied", False)
        and getattr(ctx, "reference_layout", "report") == "letter"
    )
    if not reference_letter:
        # ---- Header banner ----
        if not header_template_path:
            if getattr(ctx, "report_style", "uw") == "wa_school":
                _add_neutral_header(doc, getattr(ctx, "organization", ""))
            else:
                _add_header_with_logo(doc)

        # ---- Memo metadata ----
        # Date is intentionally blank when report_date is None — author fills
        # in the final date manually in Word before signing.
        date_text = ctx.report_date.strftime("%B %d, %Y") if ctx.report_date else ""
        for label, value in [
            ("Date", date_text),
            ("To", f"{ctx.building} Community" if getattr(ctx, "report_style", "uw") == "wa_school"
             else f"{ctx.building} Occupants"),
            ("From", getattr(ctx, "organization", "") or "Environmental Health & Safety Department"),
            ("Subject", "Water Sampling Test Results"),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(f"{label}: ")
            r.bold = True
            p.add_run(value)

    # ---- Introduction ----
    if not getattr(ctx, "reference_style_applied", False):
        _add_heading(doc, "Introduction")
    _add_md_paragraphs(doc, ctx.introduction_md)

    # ---- Testing information ----
    # Layout per UW EH&S memo style:
    #   1. Intro paragraph (sampling context + highlight legend)
    #   2. Lead regulatory context paragraph (15/10/5 ppb tiers)
    #   3. Table 1
    #   4. Detection-limit footnote
    #   5. Findings paragraph (which metals exceeded, with counts/locations)
    if not getattr(ctx, "reference_style_applied", False):
        _add_heading(doc, "Testing Information")
        doc.add_paragraph(_build_testing_info_intro(ctx))
        doc.add_paragraph(_LEAD_REGULATORY_CONTEXT)

    # ---- Table 1 ----
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Table 1: Drinking Water Fixture Sampling Results — Metal Concentrations")
    r.bold = True

    rows = _build_rows(ctx.samples, registry, ctx.action_levels,
                       ctx.analytes_shown)
    _render_results_table(doc, rows, ctx.analytes_shown)

    # Footnote
    p = doc.add_paragraph()
    r = p.add_run(
        'Values reported as "<X" indicate concentrations below the laboratory\'s '
        "detection limit of X. Highlighted cells exceed the default review levels."
    )
    r.font.size = Pt(9)
    r.italic = True

    # Findings summary — generated from the actual measurements
    if not getattr(ctx, "reference_style_applied", False):
        doc.add_paragraph(_build_findings_summary(ctx, registry))

    # ---- Actions taken ----
    if not getattr(ctx, "reference_style_applied", False):
        _add_heading(doc, "Actions Taken")
    _add_md_paragraphs(doc, ctx.actions_taken_md)

    if not getattr(ctx, "reference_style_applied", False):
        # ---- Health info (standard-template boilerplate) ----
        _add_heading(doc, "Lead Sources and Health Information")
        doc.add_paragraph(
        "People can be exposed to lead from a variety of environmental "
        "sources. Each exposure contributes to the amount of lead in the "
        "body. Some common exposure sources include:"
    )
        bullets = [
        "Dust from old, deteriorating lead paint",
        "Contaminated soil",
        "Lead dust tracked into the home from external sources, such as "
            "certain industries where lead is present",
        "Food, imported spices, cooking pots and utensils",
        "Water: It is important to reduce exposure from every source as much "
            "as possible. Lead can enter drinking water at the point-of-use "
            "(e.g., water fountains) due to lead-containing fixture "
            "components or other plumbing elements corroding. Common sources "
            "of lead in drinking water include fixture components, premise "
            "plumbing, solder, fittings, and/or lead service lines.",
        ]
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_paragraph(
        "Children six years old and younger are the most susceptible to the "
        "effects of lead. Their growing bodies absorb more lead than adults, "
        "and their brains and nervous systems are more sensitive to the "
        "damaging effects of lead. The children of women who are exposed to "
        "lead before or during pregnancy can have increased risk of these "
        "negative health effects. Adults with elevated exposure can have "
        "increased risks of heart disease, high blood pressure, kidney "
        "damage, or nervous system problems."
    )

    # ---- Questions / contacts ----
    # When the standard EH&S + Facilities pair is provided, render as a
    # single prose paragraph matching the recurring UW memo wording.
    # Otherwise fall back to a bulleted list (for one contact, three+, etc.).
        _add_heading(doc, "Questions")
        if len(ctx.contacts) == 2 and all(c.get("name") for c in ctx.contacts):
            c1, c2 = ctx.contacts
            doc.add_paragraph(_questions_prose(c1, c2))
        else:
            doc.add_paragraph(
            "If you have any questions about the water testing results or this "
            "communication, contact:"
            )
            for c in ctx.contacts:
                bits = [c.get("name", ""), c.get("title", "")]
                contact_line = " — ".join([b for b in bits if b])
                if c.get("phone"):
                    contact_line += f" — {c['phone']}"
                if c.get("email"):
                    contact_line += f", {c['email']}"
                doc.add_paragraph(contact_line, style="List Bullet")

    if ctx.notes_md:
        if not getattr(ctx, "reference_style_applied", False):
            _add_heading(doc, "Notes")
        _add_md_paragraphs(doc, ctx.notes_md)

    if isinstance(output, BytesIO):
        doc.save(output)
        output.seek(0)
        return output
    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
