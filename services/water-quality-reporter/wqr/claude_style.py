"""Create an editable communication draft from a user-supplied style sample.

The reference file controls organization, tone, headings, and reader address.
All school-specific facts are supplied separately from deterministic parsing;
Claude is explicitly forbidden from copying facts from the reference sample.
"""
from __future__ import annotations

import base64
import json
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from .claude_pdf import MODEL


TOOL_NAME = "record_style_adapted_communication"
MAX_REFERENCE_BYTES = 23 * 1024 * 1024
MAX_REFERENCE_TEXT = 60_000
MAX_DRAFT_TEXT = 12_000


class ClaudeStyleError(RuntimeError):
    """Safe error for reference-style drafting failures."""


class ClaudeStyleConfigurationError(ClaudeStyleError):
    """Raised when Claude is not configured."""


STYLE_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "layout_type": {"type": "string", "enum": ["letter", "memo", "report"]},
        "introduction_md": {"type": "string"},
        "actions_md": {"type": "string"},
        "notes_md": {"type": "string"},
        "style_summary": {"type": "string"},
    },
    "required": [
        "layout_type", "introduction_md", "actions_md", "notes_md", "style_summary"
    ],
}


_SYSTEM_PROMPT = """You adapt a school drinking-water communication to a
user-supplied reference document. Match the reference's communication genre,
section order, heading style, tone, level of detail, and reader address. Do not
copy its school name, district, dates, sample counts, test results, actions,
contacts, email addresses, phone numbers, URLs, or other case-specific facts.

Use only the LOCKED FACTS supplied by the application for the new draft. Never
change a number, threshold, date, result, school, or organization. Never claim
that remediation, shutoff, notification, consultation, or follow-up has already
happened unless it appears in confirmed_actions. If an action is not confirmed,
write it as a recommended or planned next step. Use [confirm ...] placeholders
for missing district contacts, links, or operational details.

Use school_name from LOCKED FACTS everywhere the current school is named or
addressed. Use building_names, collection_date_range, outlet counts, and lead
results from LOCKED FACTS when the reference includes corresponding details.
Do not retain any case name or case detail from the reference document.

Classify the reference layout as letter, memo, or report. The application
inserts a verified fixture-results table after introduction_md,
so do not recreate the table. Return Markdown limited to **bold**, *italic*,
blank-line paragraph breaks, and hyphen bullets. Put the reference-style title,
greeting, background, and results-summary section in introduction_md. Put the
reference-style response/action section in actions_md. Put health context,
resources, closing material, and learn-more sections in notes_md when those
elements occur in the reference. Keep the draft editable and audience-facing.
"""


def _docx_text(raw_bytes: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(raw_bytes))
    except Exception as exc:
        raise ClaudeStyleError("The Word style sample could not be read.") from exc
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                blocks.append(text)
    return "\n\n".join(blocks)[:MAX_REFERENCE_TEXT]


def _tool_payload(response: Any) -> dict:
    for block in getattr(response, "content", []) or []:
        if (
            getattr(block, "type", None) == "tool_use"
            and getattr(block, "name", None) == TOOL_NAME
            and isinstance(getattr(block, "input", None), dict)
        ):
            return block.input
    raise ClaudeStyleError("Claude returned no schema-validated style draft.")


def _clean_draft_text(value: Any, field: str, *, required: bool = False) -> str:
    if not isinstance(value, str):
        raise ClaudeStyleError(f"Claude returned an invalid {field}.")
    text = value.replace("\r\n", "\n").replace("\r", "\n").strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    if required and not text:
        raise ClaudeStyleError(f"Claude omitted {field}.")
    return text[:MAX_DRAFT_TEXT]


def _validated_draft(payload: dict) -> dict[str, str]:
    layout = payload.get("layout_type")
    if layout not in {"letter", "memo", "report"}:
        raise ClaudeStyleError("Claude returned an invalid communication layout.")
    return {
        "layout": layout,
        "intro": _clean_draft_text(
            payload.get("introduction_md"), "introduction", required=True
        ),
        "actions": _clean_draft_text(
            payload.get("actions_md"), "actions", required=True
        ),
        "notes": _clean_draft_text(payload.get("notes_md"), "notes"),
        "style_summary": _clean_draft_text(
            payload.get("style_summary"), "style summary", required=True
        )[:1000],
    }


def draft_school_communication_from_reference(
    reference_bytes: bytes,
    filename: str,
    locked_facts: dict,
    *,
    client: Any | None = None,
) -> dict[str, str]:
    """Draft communication in the reference's style using locked local facts."""
    if not reference_bytes or len(reference_bytes) > MAX_REFERENCE_BYTES:
        raise ClaudeStyleError("The report-style sample is empty or too large.")

    suffix = Path(filename).suffix.lower()
    reference_content: list[dict[str, Any]] = []
    if suffix == ".pdf":
        if not reference_bytes.startswith(b"%PDF-"):
            raise ClaudeStyleError("The report-style sample is not a valid PDF.")
        reference_content.append({
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": base64.standard_b64encode(reference_bytes).decode("ascii"),
            },
        })
    elif suffix == ".docx":
        reference_content.append({
            "type": "text",
            "text": "REFERENCE DOCUMENT TEXT:\n\n" + _docx_text(reference_bytes),
        })
    elif suffix == ".doc":
        raise ClaudeStyleError(
            "Legacy .doc style samples must be saved as .docx or PDF first."
        )
    else:
        raise ClaudeStyleError("The report-style sample must be PDF or DOCX.")

    if client is None:
        api_key = os.environ.get("CLAUDE_API_KEY")
        if not api_key:
            raise ClaudeStyleConfigurationError(
                "CLAUDE_API_KEY is not configured for style adaptation."
            )
        try:
            from anthropic import Anthropic
        except ImportError as exc:
            raise ClaudeStyleConfigurationError(
                "Claude style-adaptation dependencies are not installed."
            ) from exc
        client = Anthropic(api_key=api_key)

    reference_content.append({
        "type": "text",
        "text": (
            f"Use the uploaded reference named {Path(filename).name!r} only as a "
            "communication-style example. Draft the new communication from these "
            "LOCKED FACTS:\n"
            + json.dumps(locked_facts, ensure_ascii=False, sort_keys=True)
        ),
    })
    try:
        response = client.messages.create(
            model=MODEL,
            max_tokens=6000,
            temperature=0,
            system=_SYSTEM_PROMPT,
            tools=[{
                "name": TOOL_NAME,
                "description": "Record the editable reference-style communication draft.",
                "strict": True,
                "input_schema": STYLE_SCHEMA,
            }],
            tool_choice={
                "type": "tool",
                "name": TOOL_NAME,
                "disable_parallel_tool_use": True,
            },
            messages=[{"role": "user", "content": reference_content}],
        )
    except Exception as exc:
        raise ClaudeStyleError("Claude style adaptation is temporarily unavailable.") from exc
    return _validated_draft(_tool_payload(response))
